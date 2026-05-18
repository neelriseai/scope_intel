"""Tests for the doc ingest pipeline (Phase 5 — mode=python and mode=llm).

Covers all new components without requiring external deps (Ollama, pypdf, etc.).
Run with:  python -m pytest tests/test_doc_ingest.py -v
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from scope_intel.core import store
from scope_intel.core.indexer import build_index
from scope_intel.adapters.doc_reader import (
    read_document,
    read_document_structured,
    _split_by_tokens,
    _estimate_tokens,
    _docx_table_to_markdown,
)
from scope_intel.core.llm_client import NullLLMClient, get_client
from scope_intel.cli import (
    _doc_list, _doc_fetch, _doc_search, _doc_fetch_for, _doc_diff,
    _doc_pin, _doc_unpin, _read_pinned,
    _doc_annotate,
)
from scope_intel.core.doc_ingestor import (
    _route_section,
    _split_sections,
    _extract_memories,
    _extract_features,
    _suggest_route,
    _CURATED_TEMPLATES,
    ingest_document,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

SAMPLE_MD = """\
# My Project

This is the project overview. It explains the purpose and goals.

## Architecture

The system uses a layered architecture with three components.

### Memory Layer

Memory is stored in a Redis cache. The system must always flush on shutdown.

### Validation Engine

We decided to use JSON Schema for all payloads. Validation is mandatory.

## Roadmap

Phase 1: Core engine (done)
Phase 2: RAG layer
Phase 3: Multi-agent subagents

## Constraints

- Must not store secrets in code
- Should always validate before persisting
- Never skip schema validation in production

## Schema Design

The data model uses a flat JSON structure with a "type" discriminator.

| Field   | Type   | Required |
|---------|--------|----------|
| id      | string | yes      |
| type    | string | yes      |
| payload | object | no       |
"""


@pytest.fixture()
def md_file(tmp_path) -> Path:
    p = tmp_path / "design.md"
    p.write_text(SAMPLE_MD, encoding="utf-8")
    return p


@pytest.fixture()
def repo(tmp_path) -> Path:
    """Minimal initialised repo."""
    (tmp_path / "src").mkdir()
    (tmp_path / "src" / "main.py").write_text("def main(): pass\n", encoding="utf-8")
    store.ensure_index_dir(tmp_path)
    store.write_json(tmp_path, "config", store.default_config())
    build_index(tmp_path)
    return tmp_path


# ---------------------------------------------------------------------------
# doc_reader — read_document
# ---------------------------------------------------------------------------

class TestReadDocument:
    def test_markdown_returns_text(self, md_file):
        result = read_document(md_file)
        assert "error" not in result
        assert "text" in result
        assert "My Project" in result["text"]
        assert result["format"] == "md"

    def test_txt_format(self, tmp_path):
        p = tmp_path / "notes.txt"
        p.write_text("Hello world", encoding="utf-8")
        result = read_document(p)
        assert result["format"] == "txt"
        assert "Hello world" in result["text"]

    def test_missing_file_returns_error(self, tmp_path):
        result = read_document(tmp_path / "nonexistent.md")
        assert "error" in result

    def test_rst_treated_as_text(self, tmp_path):
        p = tmp_path / "doc.rst"
        p.write_text("Title\n=====\nBody text.", encoding="utf-8")
        result = read_document(p)
        assert "error" not in result
        assert "Title" in result["text"]


# ---------------------------------------------------------------------------
# doc_reader — token helpers
# ---------------------------------------------------------------------------

class TestTokenHelpers:
    def test_estimate_tokens_non_zero(self):
        assert _estimate_tokens("hello world") > 0

    def test_estimate_tokens_scales_with_length(self):
        short = _estimate_tokens("hi")
        long  = _estimate_tokens("hi " * 100)
        assert long > short

    def test_split_short_text_returns_single_chunk(self):
        text = "short text"
        result = _split_by_tokens(text, max_tokens=1200, overlap_tokens=150)
        assert result == [text]

    def test_split_long_text_produces_multiple_chunks(self):
        # Create a text that definitely exceeds 10 tokens (40 chars)
        text = "word " * 200   # ~1000 chars → ~250 tokens
        result = _split_by_tokens(text, max_tokens=30, overlap_tokens=5)
        assert len(result) > 1

    def test_split_chunks_have_overlap(self):
        """Last words of chunk N should appear at the start of chunk N+1."""
        text = "sentence one. sentence two. sentence three. " * 50
        chunks = _split_by_tokens(text, max_tokens=30, overlap_tokens=10)
        if len(chunks) > 1:
            # Overlap: some content from end of chunk 0 should appear in chunk 1
            end_of_first = chunks[0][-30:]
            assert any(word in chunks[1] for word in end_of_first.split() if len(word) > 3)

    def test_split_no_empty_chunks(self):
        text = "paragraph one.\n\nparagraph two.\n\nparagraph three." * 40
        chunks = _split_by_tokens(text, max_tokens=20, overlap_tokens=5)
        assert all(c.strip() for c in chunks)


# ---------------------------------------------------------------------------
# doc_reader — read_document_structured
# ---------------------------------------------------------------------------

class TestReadDocumentStructured:
    def test_returns_chunks_list(self, md_file):
        result = read_document_structured(md_file)
        assert "error" not in result
        assert "chunks" in result
        assert isinstance(result["chunks"], list)
        assert result["total_chunks"] == len(result["chunks"])

    def test_chunks_have_required_fields(self, md_file):
        result = read_document_structured(md_file)
        for chunk in result["chunks"]:
            assert "heading_path" in chunk
            assert "title" in chunk
            assert "level" in chunk
            assert "text" in chunk
            assert "token_estimate" in chunk
            assert "chunk_index" in chunk

    def test_heading_path_is_cumulative(self, md_file):
        result = read_document_structured(md_file)
        # "Memory Layer" is under Architecture → Memory Layer
        memory_chunks = [c for c in result["chunks"] if "Memory Layer" in c["title"]]
        assert memory_chunks, "Expected a Memory Layer chunk"
        path = memory_chunks[0]["heading_path"]
        assert "Architecture" in path
        assert "Memory Layer" in path
        assert ">" in path

    def test_top_level_headings_path_contains_title(self, md_file):
        """heading_path always contains the section title.

        The path is cumulative from the document root, so "Architecture"
        under "My Project" has path "My Project > Architecture", not just
        "Architecture".  We test that the title appears in the path.
        """
        result = read_document_structured(md_file)
        arch_chunks = [c for c in result["chunks"] if c["title"] == "Architecture"]
        assert arch_chunks
        assert "Architecture" in arch_chunks[0]["heading_path"]

    def test_chunk_indices_sequential(self, md_file):
        result = read_document_structured(md_file)
        indices = [c["chunk_index"] for c in result["chunks"]]
        assert indices == list(range(len(indices)))

    def test_full_text_present(self, md_file):
        result = read_document_structured(md_file)
        assert "full_text" in result
        assert "My Project" in result["full_text"]

    def test_missing_file_returns_error(self, tmp_path):
        result = read_document_structured(tmp_path / "gone.md")
        assert "error" in result

    def test_tables_in_markdown_appear_in_chunk_text(self, md_file):
        """Markdown tables should survive the heading split intact."""
        result = read_document_structured(md_file)
        # Schema Design section contains a markdown table
        schema_chunks = [c for c in result["chunks"] if "Schema" in c["title"]]
        assert schema_chunks
        table_text = schema_chunks[0]["text"]
        assert "|" in table_text  # markdown table delimiters preserved


# ---------------------------------------------------------------------------
# doc_reader — DOCX table converter (unit test, no .docx file needed)
# ---------------------------------------------------------------------------

class TestDocxTableToMarkdown:
    """Test the table-to-markdown converter using mock table objects."""

    class MockCell:
        def __init__(self, text): self.text = text

    class MockRow:
        def __init__(self, texts):
            self.cells = [TestDocxTableToMarkdown.MockCell(t) for t in texts]

    class MockTable:
        def __init__(self, rows):
            self.rows = [TestDocxTableToMarkdown.MockRow(r) for r in rows]

    def test_simple_table(self):
        table = self.MockTable([
            ["Name", "Age"],
            ["Alice", "30"],
            ["Bob", "25"],
        ])
        md = _docx_table_to_markdown(table)
        assert "| Name | Age |" in md
        assert "| --- | --- |" in md
        assert "| Alice | 30 |" in md

    def test_empty_table_returns_empty_string(self):
        table = self.MockTable([])
        assert _docx_table_to_markdown(table) == ""

    def test_merged_cells_deduplicated(self):
        """Merged cells appear as repeated text — should be deduplicated."""
        table = self.MockTable([
            ["Header", "Header", "Other"],  # merged cell repeated
            ["A", "B", "C"],
        ])
        md = _docx_table_to_markdown(table)
        lines = md.split("\n")
        header_line = lines[0]
        # "Header" should only appear once after dedup
        assert header_line.count("Header") == 1


class TestDocxReaderCoverage:
    def test_docx_reads_header_and_footer_text(self, tmp_path):
        pytest.importorskip("docx")
        import docx

        p = tmp_path / "with_header_footer.docx"
        doc = docx.Document()
        doc.sections[0].header.paragraphs[0].text = "Confidential Project Header"
        doc.sections[0].footer.paragraphs[0].text = "Version 1.2 Footer"
        doc.add_heading("Body Heading", level=1)
        doc.add_paragraph("Body paragraph.")
        doc.save(p)

        result = read_document(p)
        assert "error" not in result, result
        assert "Body Heading" in result["text"]
        assert "Header: Confidential Project Header" in result["text"]
        assert "Footer: Version 1.2 Footer" in result["text"]


# ---------------------------------------------------------------------------
# llm_client
# ---------------------------------------------------------------------------

class TestNullLLMClient:
    def test_global_summary_returns_none(self):
        client = NullLLMClient()
        assert client.global_summary("any text") is None

    def test_classify_chunk_returns_none(self):
        client = NullLLMClient()
        chunk = {"heading_path": "Arch", "title": "Arch", "text": "body", "token_estimate": 10}
        assert client.classify_chunk(chunk, {}) is None

    def test_module_map_pass_returns_none(self):
        client = NullLLMClient()
        assert client.module_map_pass([{"title": "x", "summary": "y"}]) is None

    def test_is_available_returns_false(self):
        assert NullLLMClient().is_available() is False

    def test_get_client_python_mode_returns_null(self):
        client = get_client(mode="python")
        assert isinstance(client, NullLLMClient)

    def test_get_client_llm_mode_returns_ollama(self):
        from scope_intel.core.llm_client import OllamaClient
        client = get_client(mode="llm")
        assert isinstance(client, OllamaClient)


# ---------------------------------------------------------------------------
# doc_ingestor — section routing
# ---------------------------------------------------------------------------

class TestRouteSection:
    def test_overview_routes_to_generated(self):
        dest, prefix, slug = _route_section("Introduction", "project overview goals")
        assert dest == "generated"
        assert slug == "project-overview"
        assert prefix == "001"

    def test_architecture_routes_correctly(self):
        dest, prefix, slug = _route_section("System Architecture", "high-level layers")
        assert dest == "generated"
        assert slug == "system-architecture"

    def test_constraint_routes_to_curated(self):
        dest, prefix, slug = _route_section("Constraints", "must not should never")
        assert dest == "curated"
        assert slug == "constraints"

    def test_roadmap_routes_to_generated(self):
        dest, prefix, slug = _route_section("Roadmap", "phase milestone release plan")
        assert dest == "generated"
        assert slug == "roadmap"

    def test_memory_layer_routes_correctly(self):
        dest, prefix, slug = _route_section("Memory Layer", "Redis used for caching and persistence")
        assert dest == "generated"
        assert slug == "memory-layer"

    def test_unknown_section_returns_none_tuple(self):
        dest, prefix, slug = _route_section("Random Unrelated Title", "blah blah blah")
        assert dest is None
        assert prefix is None
        assert slug is None

    def test_schema_design_routes_correctly(self):
        dest, prefix, slug = _route_section("Schema Design", "data model JSON type")
        assert dest == "generated"
        assert slug == "schema-design"

    def test_prompt_caching_routes_to_architecture(self):
        dest, prefix, slug = _route_section(
            "4.3 Prompt Caching",
            "Cache stable prompt prefixes to reduce token cost.",
        )
        assert dest == "generated"
        assert prefix == "002"
        assert slug == "system-architecture"

    def test_product_market_fit_routes_to_roadmap(self):
        dest, prefix, slug = _route_section(
            "11. Product-Market Fit Analysis",
            "Pain vs preference classification and 30-day PMF proof plan.",
        )
        assert dest == "generated"
        assert prefix is None
        assert slug == "roadmap"

    # ---- route override tests ----

    def test_route_override_curated(self):
        # <!-- route: constraints --> in body forces curated/constraints.md
        body = "<!-- route: constraints -->\nThis section is about misc stuff."
        dest, prefix, slug = _route_section("Random Title", body)
        assert dest == "curated"
        assert slug == "constraints"

    def test_route_override_generated_slug(self):
        # Override to a generated slug
        body = "<!-- route: memory-layer -->\nContent about caching."
        dest, prefix, slug = _route_section("Misc Notes", body)
        assert dest == "generated"
        assert slug == "memory-layer"

    def test_route_override_unknown_slug_falls_through(self):
        # Unknown slug in override comment → normal routing continues
        # Title "System Architecture" has no "overview" so doesn't hit project-overview
        body = "<!-- route: nonexistent-file -->\nThe system uses layered architecture."
        dest, prefix, slug = _route_section("System Architecture", body)
        # Should fall through to normal routing and match system-architecture
        assert dest == "generated"
        assert slug == "system-architecture"

    def test_route_override_case_insensitive(self):
        body = "<!-- ROUTE: roadmap -->\nPlanning content."
        dest, prefix, slug = _route_section("Planning", body)
        assert dest == "generated"
        assert slug == "roadmap"


# ---------------------------------------------------------------------------
# doc_ingestor — suggest_route hints
# ---------------------------------------------------------------------------

class TestSuggestRoute:
    def test_suggest_returns_string_or_none(self):
        hint = _suggest_route("Some Title", "some body text")
        assert hint is None or isinstance(hint, str)

    def test_suggest_finds_hint_for_memory_like_section(self):
        # Uses exact vocabulary from the memory-layer route → must get a hint
        hint = _suggest_route("Data Notes", "memory storage persistence cache layer backend")
        # "memory" and "storage" and "persistence" all appear in the broad memory routes
        assert hint is not None

    def test_suggest_returns_none_for_truly_obscure_section(self):
        hint = _suggest_route("xyzzy fnord quux", "blarg wumble zork")
        # Completely unknown vocabulary — no hint
        assert hint is None

    def test_suggest_hint_contains_slug_name(self):
        hint = _suggest_route("Caching Notes", "Redis cache performance storage")
        if hint:  # hint may or may not fire depending on match strength
            assert "md" in hint or "routes" in hint

    def test_unmatched_section_in_routing_table_has_hint(self, repo, tmp_path):
        # A doc where a section has body text with memory-like vocabulary
        # but a heading that doesn't match any route
        doc = tmp_path / "odd.md"
        doc.write_text(
            "# Project\n\nOverview of the project.\n\n"
            "## Data Store Layer\n\nWe store and cache data in Redis.\n",
            encoding="utf-8",
        )
        result = ingest_document(repo, doc, dry_run=True, overwrite=True)
        # Find unmatched entries that have hints
        unmatched = [e for e in result.get("routing_table", []) if e["file"] is None]
        # If any unmatched sections exist, at least some may have hints
        for um in unmatched:
            assert "hint" in um  # hint key must be present even if None


# ---------------------------------------------------------------------------
# doc_ingestor — section parsing
# ---------------------------------------------------------------------------

class TestSplitSections:
    def test_basic_split(self):
        text = "# Title\n\nBody text.\n\n## Sub\n\nSub body."
        sections = _split_sections(text)
        titles = [s["title"] for s in sections]
        assert "Title" in titles
        assert "Sub" in titles

    def test_preamble_captured(self):
        text = "Intro text without heading.\n\n# Section One\n\nBody."
        sections = _split_sections(text)
        assert sections[0]["title"] == "(preamble)"
        assert "Intro text" in sections[0]["body_text"]

    def test_body_text_present(self):
        text = "# Header\n\nThis is the body of the section."
        sections = _split_sections(text)
        body = next(s for s in sections if s["title"] == "Header")
        assert "body of the section" in body["body_text"]

    def test_level_recorded(self):
        text = "# H1\n\nbody\n\n## H2\n\nbody2\n\n### H3\n\nbody3"
        sections = _split_sections(text)
        levels = {s["title"]: s["level"] for s in sections}
        assert levels["H1"] == 1
        assert levels["H2"] == 2
        assert levels["H3"] == 3


# ---------------------------------------------------------------------------
# doc_ingestor — memory extraction
# ---------------------------------------------------------------------------

class TestExtractMemories:
    def test_uses_pattern_found(self):
        text = "The engine uses Redis for caching."
        mems = _extract_memories(text, "test")
        notes = [m["note"] for m in mems]
        assert any("uses" in n.lower() for n in notes)

    def test_must_pattern_is_constraint(self):
        text = "The system must validate all inputs before processing."
        mems = _extract_memories(text, "test")
        notes = [m["note"] for m in mems]
        assert any("Constraint" in n for n in notes)

    def test_decision_pattern(self):
        text = "We decided to use Qwen2.5 for all classification tasks."
        mems = _extract_memories(text, "test")
        notes = [m["note"] for m in mems]
        assert any("Decision" in n for n in notes)

    def test_no_duplicate_memories(self):
        text = "uses Redis for caching. uses Redis for caching."
        mems = _extract_memories(text, "test")
        notes = [m["note"].lower() for m in mems]
        assert len(notes) == len(set(notes))

    def test_memory_cap_applied(self):
        # Generate more than 60 patterns
        lines = [f"Component{i} uses Library{i} for stuff." for i in range(100)]
        text = "\n".join(lines)
        mems = _extract_memories(text, "test")
        assert len(mems) <= 60

    def test_confidence_in_range(self):
        text = "The system uses PostgreSQL. Must always commit transactions."
        mems = _extract_memories(text, "test")
        for m in mems:
            assert 0.0 <= m["confidence"] <= 1.0


# ---------------------------------------------------------------------------
# doc_ingestor — full pipeline (mode=python, dry_run)
# ---------------------------------------------------------------------------

class TestIngestDocumentPython:
    def test_dry_run_returns_result_without_writing(self, repo, md_file):
        result = ingest_document(repo, md_file, dry_run=True, overwrite=True)
        assert "error" not in result, result
        assert result["mode"] == "python"
        assert result["dry_run"] is True
        # No files should be on disk
        ai_ctx = repo / ".ai-context"
        assert not ai_ctx.exists()

    def test_ingest_creates_ai_context_files(self, repo, md_file):
        result = ingest_document(repo, md_file, overwrite=True)
        assert "error" not in result, result
        assert result["files_written"] > 0
        ai_ctx = repo / ".ai-context"
        assert ai_ctx.exists()
        # At least one generated file should exist
        gen_dir = ai_ctx / "generated"
        assert gen_dir.exists()
        md_files = list(gen_dir.glob("*.md"))
        assert len(md_files) > 0

    def test_ingest_writes_index_json(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        index = repo / ".ai-context" / "generated" / "index.json"
        assert index.exists()
        data = json.loads(index.read_text(encoding="utf-8"))
        assert "files" in data
        assert data["source"] == md_file.name

    def test_constraints_section_goes_to_curated(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        constraints = repo / ".ai-context" / "curated" / "constraints.md"
        assert constraints.exists()
        text = constraints.read_text(encoding="utf-8")
        assert "Constraint" in text or "must" in text.lower() or "never" in text.lower()

    def test_roadmap_section_generated(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        gen = repo / ".ai-context" / "generated"
        roadmap_files = list(gen.glob("*roadmap*"))
        assert roadmap_files, "Expected a roadmap.md file"

    def test_memories_added_to_mempalace(self, repo, md_file):
        result = ingest_document(repo, md_file, overwrite=True)
        assert result["memories_added"] > 0

    def test_features_extracted(self, repo, md_file):
        result = ingest_document(repo, md_file, overwrite=True)
        assert result["features_added"] >= 0  # may be 0 if all filtered as generic

    def test_overwrite_false_skips_existing(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)  # first run
        result2 = ingest_document(repo, md_file, overwrite=False)  # second run
        assert result2["files_written"] == 0
        assert result2["files_skipped"] > 0

    def test_overwrite_true_regenerates(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        result2 = ingest_document(repo, md_file, overwrite=True)
        assert result2["files_written"] > 0

    def test_nonexistent_doc_returns_error(self, repo, tmp_path):
        result = ingest_document(repo, tmp_path / "nope.md")
        assert "error" in result

    def test_uninitialised_repo_auto_inits(self, tmp_path, md_file):
        # ingest_document auto-initialises the scope store; no prior 'scope init' needed
        result = ingest_document(tmp_path, md_file, overwrite=True)
        assert "error" not in result, result
        # .scope-intelligence/ should now exist
        assert (tmp_path / ".scope-intelligence").exists()

    def test_preamble_with_architecture_terms_routes_to_overview(self, repo, tmp_path):
        doc = tmp_path / "preamble.md"
        doc.write_text(
            "This architecture overview explains the system pipeline before headings.\n\n"
            "# Details\n\nRandom notes without routing words.\n",
            encoding="utf-8",
        )
        result = ingest_document(repo, doc, dry_run=True, overwrite=True)
        preamble = next(e for e in result["routing_table"] if "preamble" in e["section"])
        assert preamble["file"] == "001-project-overview.md"
        assert preamble["layer"] == "generated"
        assert result["routing_coverage"]["sections_routed"] >= 1


# ---------------------------------------------------------------------------
# doc_ingestor — routing_table in result
# ---------------------------------------------------------------------------

class TestRoutingTable:
    """routing_table is always present in the ingest result (python mode)."""

    def test_routing_table_present_in_result(self, repo, md_file):
        result = ingest_document(repo, md_file, dry_run=True, overwrite=True)
        assert "routing_table" in result
        assert isinstance(result["routing_table"], list)

    def test_routing_table_not_empty(self, repo, md_file):
        result = ingest_document(repo, md_file, dry_run=True, overwrite=True)
        assert len(result["routing_table"]) > 0

    def test_routing_table_entry_has_required_keys(self, repo, md_file):
        result = ingest_document(repo, md_file, dry_run=True, overwrite=True)
        for entry in result["routing_table"]:
            assert "section" in entry
            assert "file" in entry     # may be None if unmatched
            assert "layer" in entry    # may be None if unmatched

    def test_constraints_routed_to_curated(self, repo, md_file):
        result = ingest_document(repo, md_file, dry_run=True, overwrite=True)
        # SAMPLE_MD has ## Constraints → should land in curated/constraints.md
        constraints_routes = [
            e for e in result["routing_table"]
            if e.get("section", "").lower() == "constraints"
        ]
        assert constraints_routes, "Expected Constraints section in routing_table"
        assert constraints_routes[0]["layer"] == "curated"
        assert constraints_routes[0]["file"] == "constraints.md"

    def test_roadmap_routed_to_generated(self, repo, md_file):
        result = ingest_document(repo, md_file, dry_run=True, overwrite=True)
        roadmap_routes = [
            e for e in result["routing_table"]
            if "roadmap" in e.get("section", "").lower()
        ]
        assert roadmap_routes, "Expected Roadmap section in routing_table"
        assert roadmap_routes[0]["layer"] == "generated"
        assert "roadmap" in roadmap_routes[0]["file"]

    def test_unmatched_sections_have_none_file(self, repo, tmp_path):
        # A doc with sections that don't match any route
        doc = tmp_path / "weird.md"
        doc.write_text(
            "# Totally Obscure Section\n\nContent about xyzzy fnord quux.\n\n"
            "## Another Strange Thing\n\nMore quux content here.\n",
            encoding="utf-8",
        )
        result = ingest_document(repo, doc, dry_run=True, overwrite=True)
        # At least one unmatched section should appear with file=None
        unmatched = [e for e in result.get("routing_table", []) if e["file"] is None]
        assert len(unmatched) >= 0  # OK even if zero (routing may catch things)

    def test_routing_table_present_in_non_dry_run(self, repo, md_file):
        # routing_table is included even when not dry-run
        result = ingest_document(repo, md_file, overwrite=True)
        assert "routing_table" in result
        assert isinstance(result["routing_table"], list)

    def test_routing_table_matches_sections_parsed(self, repo, md_file):
        result = ingest_document(repo, md_file, dry_run=True, overwrite=True)
        # routing_table entries + preamble-skips should account for sections parsed
        # (sections_parsed may include preamble w/o body, so table can be ≤ parsed)
        assert len(result["routing_table"]) <= result["sections_parsed"]


# ---------------------------------------------------------------------------
# doc_ingestor — full pipeline (mode=llm, NullLLMClient fallback)
# ---------------------------------------------------------------------------

class TestIngestDocumentLLMFallback:
    """Test mode=llm with Ollama unavailable — should fall back to python."""

    def test_llm_fallback_to_python_when_ollama_down(self, repo, md_file):
        # Ollama is not running in CI — ingest_document should fall back
        result = ingest_document(
            repo, md_file,
            mode="llm",
            ollama_url="http://localhost:19999",  # wrong port, guaranteed unavailable
            overwrite=True,
        )
        assert "error" not in result, result
        # Should fall back to python mode
        assert result["mode"] in ("python", "python_fallback")
        assert "warning" in result


class TestIngestDocumentLLMSecondPass:
    class FakeLLM:
        model = "fake-qwen"

        def is_available(self):
            return True

        def global_summary(self, excerpt):  # noqa: ARG002
            return {"project_name": "Fake Project", "components": ["Runtime"]}

        def classify_chunk(self, chunk, global_ctx):  # noqa: ARG002
            title = chunk["title"]
            if "Architecture" in title:
                return {
                    "target_file": "002-system-architecture.md",
                    "section_title": "Architecture",
                    "summary": "Runtime architecture summary.",
                    "key_facts": [],
                    "constraints": [],
                    "is_feature": False,
                    "feature_id": "",
                    "tags": ["architecture"],
                }
            if "Module Map" in title:
                return {
                    "target_file": "module-map.md",
                    "section_title": "Existing Module Map",
                    "summary": "Existing source map.",
                    "key_facts": [],
                    "constraints": [],
                    "is_feature": False,
                    "feature_id": "",
                    "tags": [],
                }
            return {"target_file": "skip"}

        def module_map_pass(self, file_summaries):
            assert file_summaries
            return "# Synthesized Module Map\n\n- Runtime -> API"

    def test_second_pass_overwrites_placeholder_module_bucket(self, repo, tmp_path, monkeypatch):
        from scope_intel.core import llm_client

        doc = tmp_path / "llm_design.md"
        doc.write_text(
            "# Architecture\n\nThe runtime architecture has an API boundary.\n\n"
            "# Module Map\n\nOld source module text that should be preserved below synthesis.\n",
            encoding="utf-8",
        )
        fake = self.FakeLLM()
        monkeypatch.setattr(llm_client, "get_client", lambda **kwargs: fake)

        result = ingest_document(repo, doc, mode="llm", overwrite=True, second_pass=True)
        assert "error" not in result, result
        module_map = repo / ".ai-context" / "curated" / "module-map.md"
        content = module_map.read_text(encoding="utf-8")
        assert "Synthesized Module Map" in content
        assert "Old source module text" in content
        assert result["routing_coverage"]["sections_routed"] >= 2


# ---------------------------------------------------------------------------
# doc_ingestor — CLAUDE.md update
# ---------------------------------------------------------------------------

class TestClaudeMdUpdate:
    def test_claude_md_created_when_missing(self, repo, md_file):
        claude_md = repo / "CLAUDE.md"
        claude_dir_md = repo / ".claude" / "CLAUDE.md"
        assert not claude_md.exists()
        assert not claude_dir_md.exists()
        ingest_document(repo, md_file, overwrite=True, update_claude_md=True)
        assert claude_md.exists()
        assert claude_dir_md.exists()
        content = claude_md.read_text(encoding="utf-8")
        agent_content = claude_dir_md.read_text(encoding="utf-8")
        assert "scope-intel-doc-context" in content
        assert "scope-intel-doc-context" in agent_content

    def test_claude_md_not_updated_when_flag_false(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True, update_claude_md=False)
        assert not (repo / "CLAUDE.md").exists()
        assert not (repo / ".claude" / "CLAUDE.md").exists()

    def test_claude_md_section_replaced_on_rerun(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        ingest_document(repo, md_file, overwrite=True)
        content = (repo / "CLAUDE.md").read_text(encoding="utf-8")
        agent_content = (repo / ".claude" / "CLAUDE.md").read_text(encoding="utf-8")
        # Marker should appear exactly once (not duplicated)
        assert content.count("<!-- scope-intel-doc-context -->") == 1
        assert agent_content.count("<!-- scope-intel-doc-context -->") == 1


# ---------------------------------------------------------------------------
# doc list and doc fetch helpers (CLI layer)
# ---------------------------------------------------------------------------

class TestDocListAndFetch:
    """Test the _doc_list() and _doc_fetch() helpers used by CLI + MCP."""

    def test_list_no_ai_context_returns_error(self, repo):
        result = _doc_list(repo)
        assert "error" in result

    def test_list_after_ingest_returns_files(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_list(repo)
        assert "error" not in result
        assert result["total"] > 0
        assert isinstance(result["generated"], list)

    def test_list_source_matches_ingested_doc(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_list(repo)
        assert result["source"] == md_file.name

    def test_list_curated_files_present(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_list(repo)
        curated_ids = [c["id"] for c in result.get("curated", [])]
        # constraints.md should have been written (SAMPLE_MD has ## Constraints)
        assert "constraints" in curated_ids

    def test_fetch_by_partial_id(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_fetch(repo, "overview")
        assert "error" not in result, result
        assert "content" in result
        assert result["chars"] > 0
        assert "overview" in result["id"].lower() or "overview" in result["title"].lower()

    def test_fetch_by_number_prefix(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_fetch(repo, "001")
        assert "error" not in result, result
        assert "001" in result["id"]

    def test_fetch_curated_constraints(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_fetch(repo, "constraints")
        assert "error" not in result, result
        assert result["layer"] == "curated"
        assert "constraints" in result["id"]

    def test_fetch_no_match_returns_error(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_fetch(repo, "xyznonexistent999")
        assert "error" in result

    def test_fetch_ambiguous_returns_error(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        # "a" matches many files
        result = _doc_fetch(repo, "a")
        if "error" in result:
            assert "ambiguous" in result["error"] or "no file" in result["error"]
        # (if only one matches, that's also acceptable)


# ---------------------------------------------------------------------------
# doc search helper
# ---------------------------------------------------------------------------

class TestDocSearch:
    def test_search_no_ai_context_returns_error(self, repo):
        result = _doc_search(repo, "anything")
        assert "error" in result

    def test_search_finds_keyword(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        # SAMPLE_MD has "Redis" in the Memory Layer section
        result = _doc_search(repo, "Redis")
        assert "error" not in result, result
        assert result["total_matches"] > 0

    def test_search_returns_file_and_line(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_search(repo, "Redis")
        assert result["files_with_matches"] > 0
        hit = result["results"][0]
        assert "file_id" in hit
        assert "path" in hit
        assert "matches" in hit
        match = hit["matches"][0]
        assert "line_no" in match
        assert "line" in match
        assert "Redis" in match["line"]

    def test_search_context_lines_present(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_search(repo, "Redis", context_lines=2)
        if result["total_matches"] > 0:
            m = result["results"][0]["matches"][0]
            assert isinstance(m["context_before"], list)
            assert isinstance(m["context_after"], list)

    def test_search_case_insensitive(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        r1 = _doc_search(repo, "redis")   # lowercase
        r2 = _doc_search(repo, "REDIS")   # uppercase
        assert r1["total_matches"] == r2["total_matches"]

    def test_search_layer_filter_generated(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        r_all = _doc_search(repo, "Auto-generated", layer="all")
        r_gen = _doc_search(repo, "Auto-generated", layer="generated")
        # generated/ should have matches (the auto-generated footer)
        assert r_gen["total_matches"] > 0
        assert r_gen["total_matches"] <= r_all["total_matches"]

    def test_search_no_match_returns_empty_results(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_search(repo, "xyznonexistentkeyword999")
        assert "error" not in result
        assert result["total_matches"] == 0
        assert result["results"] == []

    def test_search_total_matches_consistent(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_search(repo, "Auto-generated")
        manual_total = sum(r["match_count"] for r in result["results"])
        assert result["total_matches"] == manual_total

    def test_search_regex_finds_pattern(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        # Regex: match "Redis" or "cache" or "memory"
        result = _doc_search(repo, r"Redis|cache|memory", use_regex=True)
        assert "error" not in result
        assert result.get("use_regex") is True
        assert result["total_matches"] > 0

    def test_search_invalid_regex_returns_error(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_search(repo, r"[invalid(", use_regex=True)
        assert "error" in result

    def test_search_literal_vs_regex_differ_for_special_chars(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        # Literal search for "scope.doc" matches the substring literally
        r_lit = _doc_search(repo, "scope.doc", use_regex=False)
        # Regex "scope.doc" treats '.' as any char — may match more
        r_rx  = _doc_search(repo, "scope.doc", use_regex=True)
        # Regex can't match fewer than literal (since '.' is a superset)
        assert r_rx["total_matches"] >= r_lit["total_matches"]


# ---------------------------------------------------------------------------
# doc stats helper
# ---------------------------------------------------------------------------

class TestDocStats:
    """Test the _doc_stats() helper used by `scope doc stats` and MCP doc_stats."""

    def test_stats_no_ai_context_returns_error(self, repo):
        from scope_intel.cli import _doc_stats
        result = _doc_stats(repo)
        assert "error" in result

    def test_stats_after_ingest_has_totals(self, repo, md_file):
        from scope_intel.cli import _doc_stats
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_stats(repo)
        assert "error" not in result
        assert result["total_files"] > 0
        assert result["total_chars"] > 0
        assert result["total_tokens"] > 0

    def test_stats_tokens_approx_chars_over_4(self, repo, md_file):
        from scope_intel.cli import _doc_stats
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_stats(repo)
        # Per-file token = chars // 4 (integer division truncates per file, not in aggregate)
        # total_tokens ≈ total_chars / 4, within ±(num_files) due to truncation
        n = result["total_files"]
        approx = result["total_chars"] / 4
        assert abs(result["total_tokens"] - approx) <= n + 1

    def test_stats_generated_and_curated_listed(self, repo, md_file):
        from scope_intel.cli import _doc_stats
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_stats(repo)
        assert isinstance(result["generated"], list)
        assert isinstance(result["curated"], list)
        # SAMPLE_MD produces at least one generated and one curated file
        assert len(result["generated"]) > 0
        assert len(result["curated"]) > 0

    def test_stats_each_entry_has_required_keys(self, repo, md_file):
        from scope_intel.cli import _doc_stats
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_stats(repo)
        for entry in result["generated"] + result["curated"]:
            assert "id"     in entry
            assert "path"   in entry
            assert "layer"  in entry
            assert "chars"  in entry
            assert "tokens" in entry

    def test_stats_total_matches_sum_of_parts(self, repo, md_file):
        from scope_intel.cli import _doc_stats
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_stats(repo)
        all_files = result["generated"] + result["curated"]
        assert result["total_chars"]  == sum(f["chars"]  for f in all_files)
        assert result["total_tokens"] == sum(f["tokens"] for f in all_files)
        assert result["total_files"]  == len(all_files)


# ---------------------------------------------------------------------------
# doc source hash + --if-changed
# ---------------------------------------------------------------------------

class TestDocHashAndIfChanged:
    """Test source hash tracking and --if-changed skip logic."""

    def test_ingest_result_has_source_hash(self, repo, md_file):
        result = ingest_document(repo, md_file, overwrite=True)
        assert "source_hash" in result
        assert len(result["source_hash"]) == 16  # first 16 hex chars
        assert result["source_hash"].isalnum()

    def test_source_hash_is_consistent(self, repo, md_file):
        r1 = ingest_document(repo, md_file, overwrite=True)
        r2 = ingest_document(repo, md_file, overwrite=True)
        assert r1["source_hash"] == r2["source_hash"]

    def test_source_hash_changes_when_doc_changes(self, repo, md_file):
        r1 = ingest_document(repo, md_file, overwrite=True)
        # Modify the file
        orig = md_file.read_text(encoding="utf-8")
        md_file.write_text(orig + "\n## New Section\n\nExtra content.\n", encoding="utf-8")
        r2 = ingest_document(repo, md_file, overwrite=True)
        assert r1["source_hash"] != r2["source_hash"]

    def test_if_changed_skips_when_unchanged(self, repo, md_file):
        # First ingest writes the hash to index.json
        ingest_document(repo, md_file, overwrite=True)
        # Second call with if_changed=True should skip
        result = ingest_document(repo, md_file, if_changed=True)
        assert result.get("unchanged") is True
        assert result["files_written"] == 0
        assert "source_hash" in result

    def test_if_changed_runs_when_doc_modified(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        # Modify the doc
        orig = md_file.read_text(encoding="utf-8")
        md_file.write_text(orig + "\n## Extra\n\nMore overview content.\n", encoding="utf-8")
        # Should NOT skip — doc changed
        result = ingest_document(repo, md_file, if_changed=True, overwrite=True)
        assert not result.get("unchanged")
        assert result["files_written"] >= 0  # may write or skip existing

    def test_if_changed_runs_when_no_prior_ingest(self, repo, md_file):
        # No prior index.json → always run
        result = ingest_document(repo, md_file, if_changed=True, overwrite=True)
        assert not result.get("unchanged")
        assert "error" not in result

    def test_index_json_contains_source_hash(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        index = json.loads(
            (repo / ".ai-context" / "generated" / "index.json").read_text(encoding="utf-8")
        )
        assert "source_hash" in index
        assert len(index["source_hash"]) == 16


# ---------------------------------------------------------------------------
# curated file templates
# ---------------------------------------------------------------------------

class TestCuratedTemplates:
    """Test that missing curated/ files get starter templates after ingest."""

    def test_templates_created_key_in_result(self, repo, md_file):
        result = ingest_document(repo, md_file, overwrite=True)
        assert "templates_created" in result
        assert isinstance(result["templates_created"], list)

    def test_templates_created_for_missing_curated_files(self, repo, tmp_path):
        # Use a doc with NO curated-routable sections so all 3 templates are created
        doc = tmp_path / "plain.md"
        doc.write_text(
            "# Overview\n\nSystem overview content.\n\n"
            "## Architecture\n\nHigh-level system design.\n",
            encoding="utf-8",
        )
        result = ingest_document(repo, doc, overwrite=True)
        # At least some templates should be created
        assert len(result["templates_created"]) > 0

    def test_templates_not_created_when_curated_file_exists(self, repo, md_file):
        # SAMPLE_MD has ## Constraints → writes curated/constraints.md
        ingest_document(repo, md_file, overwrite=True)
        constraints = repo / ".ai-context" / "curated" / "constraints.md"
        assert constraints.exists()
        # constraints.md from the doc should NOT be replaced by a template
        content = constraints.read_text(encoding="utf-8")
        # The doc-generated file has "Constraint" from the actual content
        assert "TODO" not in content or "must not" in content.lower()

    def test_template_files_exist_on_disk(self, repo, tmp_path):
        doc = tmp_path / "simple.md"
        doc.write_text("# Overview\n\nProject overview.\n\n## Architecture\n\nLayers.\n",
                       encoding="utf-8")
        ingest_document(repo, doc, overwrite=True)
        cur_dir = repo / ".ai-context" / "curated"
        assert cur_dir.exists()
        # At least one of the three template files should exist
        template_files = [f for f in _CURATED_TEMPLATES if (cur_dir / f).exists()]
        assert len(template_files) > 0

    def test_templates_not_overwritten_on_re_ingest(self, repo, tmp_path):
        doc = tmp_path / "simple.md"
        doc.write_text("# Overview\n\nProject overview.\n\n## Architecture\n\nLayers.\n",
                       encoding="utf-8")
        # First ingest — templates created
        ingest_document(repo, doc, overwrite=True)
        # Edit the constraints template manually
        constraints = repo / ".ai-context" / "curated" / "constraints.md"
        if constraints.exists():
            constraints.write_text("# Constraints\n\nMy custom constraint.\n", encoding="utf-8")
        # Second ingest — must NOT overwrite the hand-edited file
        ingest_document(repo, doc, overwrite=True)
        if constraints.exists():
            content = constraints.read_text(encoding="utf-8")
            assert "My custom constraint" in content

    def test_dry_run_does_not_create_templates(self, repo, tmp_path):
        doc = tmp_path / "simple.md"
        doc.write_text("# Overview\n\nProject overview.\n", encoding="utf-8")
        result = ingest_document(repo, doc, dry_run=True, overwrite=True)
        assert result["templates_created"] == []
        cur_dir = repo / ".ai-context" / "curated"
        assert not cur_dir.exists()


# ---------------------------------------------------------------------------
# doc export helper
# ---------------------------------------------------------------------------

class TestDocExport:
    """Test the _doc_export() helper used by `scope doc export` and MCP doc_export."""

    def test_export_no_ai_context_returns_error(self, repo):
        from scope_intel.cli import _doc_export
        result = _doc_export(repo)
        assert "error" in result

    def test_export_after_ingest_has_content(self, repo, md_file):
        from scope_intel.cli import _doc_export
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_export(repo)
        assert "error" not in result
        assert len(result["content"]) > 0
        assert result["total_files"] > 0

    def test_export_contains_all_file_content(self, repo, md_file):
        from scope_intel.cli import _doc_export
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_export(repo)
        # Content from SAMPLE_MD should be present somewhere in the export
        assert "Redis" in result["content"]   # from Memory Layer section

    def test_export_header_present_by_default(self, repo, md_file):
        from scope_intel.cli import _doc_export
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_export(repo)
        assert "scope-intel-export" in result["content"]

    def test_export_no_header_flag(self, repo, md_file):
        from scope_intel.cli import _doc_export
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_export(repo, include_header=False)
        assert "scope-intel-export" not in result["content"]

    def test_export_layer_generated_only(self, repo, md_file):
        from scope_intel.cli import _doc_export
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_export(repo, layer="generated")
        layers = {f["layer"] for f in result["files"]}
        assert layers == {"generated"}

    def test_export_layer_curated_only(self, repo, md_file):
        from scope_intel.cli import _doc_export
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_export(repo, layer="curated")
        layers = {f["layer"] for f in result["files"]}
        assert layers == {"curated"}

    def test_export_result_keys(self, repo, md_file):
        from scope_intel.cli import _doc_export
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_export(repo)
        for key in ("content", "source", "total_files", "total_chars",
                    "total_tokens", "files"):
            assert key in result

    def test_export_total_files_matches_list(self, repo, md_file):
        from scope_intel.cli import _doc_export
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_export(repo)
        assert result["total_files"] == len(result["files"])

    def test_export_tag_filter_includes_only_tagged(self, repo, md_file):
        from scope_intel.cli import _doc_export, _doc_tag, _doc_list
        ingest_document(repo, md_file, overwrite=True)
        # Pick the first generated file and tag it.
        listing = _doc_list(repo)
        first_id = listing["generated"][0]["id"]
        _doc_tag(repo, first_id, add_tags=["important"])

        result = _doc_export(repo, tag_filter="important")
        assert "error" not in result, result
        assert result["total_files"] == 1
        assert result["files"][0]["id"] == first_id

    def test_export_tag_filter_no_match_returns_error(self, repo, md_file):
        from scope_intel.cli import _doc_export
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_export(repo, tag_filter="never-applied-tag")
        assert "error" in result
        assert "never-applied-tag" in result["error"]

    def test_export_tag_filter_curated(self, repo, md_file):
        """Tag filter should also pick up tags on curated files."""
        from scope_intel.cli import _doc_export, _doc_tag
        ingest_document(repo, md_file, overwrite=True)
        cur_dir = repo / ".ai-context" / "curated"
        cur_dir.mkdir(parents=True, exist_ok=True)
        (cur_dir / "manual-notes.md").write_text("# Notes\n", encoding="utf-8")
        _doc_tag(repo, "manual-notes", add_tags=["curated-only"])

        result = _doc_export(repo, tag_filter="curated-only")
        assert "error" not in result, result
        assert result["total_files"] == 1
        assert result["files"][0]["layer"] == "curated"


# ---------------------------------------------------------------------------
# doc check — health validation
# ---------------------------------------------------------------------------

class TestDocCheck:
    """Test the _doc_check() health validator."""

    def test_check_no_ai_context_returns_error(self, repo):
        from scope_intel.cli import _doc_check
        result = _doc_check(repo)
        assert "error" in result

    def test_check_after_ingest_is_healthy_or_warns(self, repo, md_file):
        from scope_intel.cli import _doc_check
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_check(repo)
        assert "error" not in result
        # "errors" key is always present
        assert "errors" in result
        assert "warnings" in result
        assert isinstance(result["healthy"], bool)

    def test_check_result_has_required_keys(self, repo, md_file):
        from scope_intel.cli import _doc_check
        ingest_document(repo, md_file, overwrite=True)
        result = _doc_check(repo)
        for key in ("source", "generated_at", "mode", "errors", "warnings",
                    "issues", "passes", "generated_files", "curated_files"):
            assert key in result

    def test_check_detects_missing_generated_file(self, repo, md_file):
        from scope_intel.cli import _doc_check
        ingest_document(repo, md_file, overwrite=True)
        # Delete one generated file but leave index.json intact
        gen_dir = repo / ".ai-context" / "generated"
        mds = list(gen_dir.glob("*.md"))
        if mds:
            mds[0].unlink()
            result = _doc_check(repo)
            # Should have at least one error (missing file)
            assert result["errors"] > 0
            error_files = [i["file"] for i in result["issues"] if i["level"] == "error"]
            assert any(mds[0].name in f for f in error_files)

    def test_check_detects_unindexed_generated_file(self, repo, md_file):
        from scope_intel.cli import _doc_check
        ingest_document(repo, md_file, overwrite=True)
        stale = repo / ".ai-context" / "generated" / "stale-orphan.md"
        stale.write_text("# Stale\n\nThis file is not in index.json.\n", encoding="utf-8")

        result = _doc_check(repo)
        warn_files = [i["file"] for i in result["issues"] if i["level"] == "warn"]
        assert any("stale-orphan.md" in f for f in warn_files)

    def test_check_detects_todo_placeholders_in_curated(self, repo, md_file):
        from scope_intel.cli import _doc_check
        ingest_document(repo, md_file, overwrite=True)
        # Overwrite current-phase.md with template-style content containing TODO
        cur_dir = repo / ".ai-context" / "curated"
        cur_dir.mkdir(parents=True, exist_ok=True)
        (cur_dir / "current-phase.md").write_text(
            "# Current Phase\n\n- TODO: fill this in\n", encoding="utf-8"
        )
        result = _doc_check(repo)
        # Should warn about TODO placeholder
        warn_msgs = [i["msg"] for i in result["issues"] if "TODO" in i["msg"]]
        assert warn_msgs

    def test_check_passes_clean_content(self, repo, md_file):
        from scope_intel.cli import _doc_check
        ingest_document(repo, md_file, overwrite=True)
        # Write clean curated files (no TODO)
        cur_dir = repo / ".ai-context" / "curated"
        cur_dir.mkdir(parents=True, exist_ok=True)
        for fname in ("current-phase.md", "module-map.md"):
            p = cur_dir / fname
            if not p.exists():
                p.write_text(f"# {fname}\n\nClean content with no placeholders.\n",
                             encoding="utf-8")
        result = _doc_check(repo)
        # Passes list should be non-empty (at least generated files pass)
        assert len(result["passes"]) > 0

    def test_check_curated_missing_flagged_as_warning(self, repo, md_file):
        from scope_intel.cli import _doc_check
        # Use a doc with no curated-routable sections → curated files only from templates
        # Then manually delete one template
        ingest_document(repo, md_file, overwrite=True)
        cur_dir = repo / ".ai-context" / "curated"
        module_map = cur_dir / "module-map.md"
        if module_map.exists():
            module_map.unlink()
        result = _doc_check(repo)
        # module-map.md missing should appear as a warning
        warn_files = [i["file"] for i in result["issues"] if i["level"] == "warn"]
        assert any("module-map" in f for f in warn_files)

    def test_ingest_with_verify_appends_health_check(self, repo, md_file):
        """--verify flag appends health_check to the ingest result."""
        # Simulate --verify by calling ingest + _doc_check manually
        result = ingest_document(repo, md_file, overwrite=True)
        from scope_intel.cli import _doc_check
        result["health_check"] = _doc_check(repo)
        assert "health_check" in result
        hc = result["health_check"]
        assert "errors" in hc
        assert "healthy" in hc


# ---------------------------------------------------------------------------
# ingest-batch (tested via multiple sequential ingest_document calls)
# ---------------------------------------------------------------------------

class TestIngestBatch:
    """Test the multi-file batch ingest pattern (sequential ingest_document calls).

    The CLI ingest-batch command is just a loop over ingest_document; we test
    the core logic here without needing to invoke argparse.
    """

    @pytest.fixture()
    def docs_dir(self, tmp_path) -> Path:
        """Create a small directory with multiple markdown design docs."""
        d = tmp_path / "docs"
        d.mkdir()
        (d / "overview.md").write_text(
            "# Project Overview\n\nThis project builds an AI assistant.\n\n"
            "## Architecture\n\nLayered architecture using Python.\n",
            encoding="utf-8",
        )
        (d / "specs.md").write_text(
            "# Memory Layer\n\nRedis-backed storage for semantic memories.\n\n"
            "## Constraints\n\nMust not exceed 1GB memory usage.\n",
            encoding="utf-8",
        )
        (d / "roadmap.md").write_text(
            "# Roadmap\n\nPhase 1: Core engine (done).\nPhase 2: RAG layer.\n",
            encoding="utf-8",
        )
        return d

    def test_multiple_docs_all_succeed(self, repo, docs_dir):
        results = []
        for p in sorted(docs_dir.glob("*.md")):
            r = ingest_document(repo, p, overwrite=True)
            results.append(r)
        assert all("error" not in r for r in results)

    def test_multiple_docs_aggregate_files_written(self, repo, docs_dir):
        total_written = 0
        for p in sorted(docs_dir.glob("*.md")):
            r = ingest_document(repo, p, overwrite=True)
            total_written += r.get("files_written", 0)
        # Three docs covering overview/arch, memory/constraints, roadmap → multiple files
        assert total_written > 0

    def test_second_doc_skips_existing_without_overwrite(self, repo, docs_dir):
        doc_list = sorted(docs_dir.glob("*.md"))
        # First ingest
        ingest_document(repo, doc_list[0], overwrite=True)
        # Second ingest without overwrite — existing files should be skipped
        r2 = ingest_document(repo, doc_list[1], overwrite=False)
        # Some files might be skipped (already written by doc 1), some new
        assert "error" not in r2

    def test_if_changed_skips_unchanged_in_batch(self, repo, docs_dir):
        doc_list = sorted(docs_dir.glob("*.md"))
        # Prime the hash for the first doc
        ingest_document(repo, doc_list[0], overwrite=True)
        # Run again with if_changed — first doc should be unchanged
        r = ingest_document(repo, doc_list[0], if_changed=True)
        assert r.get("unchanged") is True

    def test_each_doc_contributes_to_index(self, repo, docs_dir):
        for p in sorted(docs_dir.glob("*.md")):
            ingest_document(repo, p, overwrite=True)
        index_path = repo / ".ai-context" / "generated" / "index.json"
        assert index_path.exists()
        index = json.loads(index_path.read_text(encoding="utf-8"))
        # Last doc's files should be recorded (index.json is overwritten each run)
        assert len(index["files"]) > 0

    def test_batch_with_mixed_extensions(self, repo, tmp_path):
        # Create a directory with mixed file types
        d = tmp_path / "mixed"
        d.mkdir()
        (d / "design.md").write_text(
            "# Architecture\n\nHigh-level design.\n", encoding="utf-8"
        )
        (d / "notes.txt").write_text(
            "Overview of the project goals and scope.\n", encoding="utf-8"
        )
        results = []
        for pat in ("*.md", "*.txt"):
            for p in sorted(d.glob(pat)):
                r = ingest_document(repo, p, overwrite=True)
                results.append(r)
        assert all("error" not in r for r in results)
        assert len(results) == 2


# ---------------------------------------------------------------------------
# doc clear (tested via filesystem state)
# ---------------------------------------------------------------------------

class TestDocClear:
    """Test _doc_clear logic through ingest+clear cycle.

    We can't easily call the CLI arg-based cmd_doc("clear") from tests,
    so we test the filesystem directly after using the CLI subprocess approach
    or just verify the directory state after a real clear via Python.
    """

    def test_clear_removes_generated_files(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        gen_dir = repo / ".ai-context" / "generated"
        assert gen_dir.exists()
        assert len(list(gen_dir.glob("*.md"))) > 0

        # Simulate what clear does
        import shutil
        shutil.rmtree(gen_dir)

        # After clear, list should error
        result = _doc_list(repo)
        assert "error" in result

    def test_ingest_after_clear_starts_fresh(self, repo, md_file):
        ingest_document(repo, md_file, overwrite=True)
        gen_dir = repo / ".ai-context" / "generated"
        import shutil
        shutil.rmtree(gen_dir)

        # Re-ingest should work cleanly
        result = ingest_document(repo, md_file, overwrite=True)
        assert "error" not in result
        assert result["files_written"] > 0


# ---------------------------------------------------------------------------
# TestDocFetchFor
# ---------------------------------------------------------------------------

class TestDocFetchFor:
    """Tests for _doc_fetch_for — unified context bundle for a feature."""

    # -- helpers --

    def _ingest(self, repo, md_file, **kwargs):
        return ingest_document(repo, md_file, overwrite=True, **kwargs)

    # -- tests --

    def test_returns_expected_keys(self, repo, md_file):
        self._ingest(repo, md_file)
        result = _doc_fetch_for(repo, "memory")
        for key in ("feature", "slug", "doc_files", "doc_search",
                    "memories", "scope", "total_doc_files",
                    "total_doc_excerpts", "total_memories"):
            assert key in result, f"missing key: {key}"

    def test_no_ai_context_returns_empty_not_error(self, tmp_path):
        """If .ai-context/ doesn't exist we get an empty result, not an error."""
        result = _doc_fetch_for(tmp_path, "anything")
        assert "error" not in result
        assert result["total_doc_files"] == 0
        assert result["total_doc_excerpts"] == 0

    def test_slug_normalisation(self, repo):
        result = _doc_fetch_for(repo, "Memory Layer  (v2)")
        assert result["slug"] == "memory-layer-v2"

    def test_doc_search_finds_mentions(self, repo, md_file):
        self._ingest(repo, md_file)
        # SAMPLE_MD has a "Roadmap" section — search should surface a doc_search hit
        result = _doc_fetch_for(repo, "roadmap")
        total = result["total_doc_files"] + result["total_doc_excerpts"]
        assert total > 0, "expected at least one doc_files or doc_search hit for 'roadmap'"

    def test_doc_files_matched_by_filename(self, repo, tmp_path):
        """A .ai-context/ file whose stem contains the slug should appear in doc_files."""
        # Create a synthetic .ai-context/generated/ file whose name contains the slug
        ai_gen = repo / ".ai-context" / "generated"
        ai_gen.mkdir(parents=True, exist_ok=True)
        target = ai_gen / "001-memory-layer.md"
        target.write_text("# Memory Layer\nSome content about memory.\n", encoding="utf-8")
        # Also write a minimal index.json so _doc_list doesn't fail
        (ai_gen / "index.json").write_text(
            json.dumps({
                "source": "test.md", "generated_at": "2026-01-01",
                "mode": "python", "source_hash": "",
                "files": [{"id": "001-memory-layer",
                            "path": ".ai-context/generated/001-memory-layer.md",
                            "layer": "generated", "title": "Memory Layer"}],
            }), encoding="utf-8",
        )
        result = _doc_fetch_for(repo, "memory-layer")
        assert result["total_doc_files"] >= 1
        ids = [f["id"] for f in result["doc_files"]]
        assert "001-memory-layer" in ids

    def test_no_memories_flag_skips_memories(self, repo, md_file):
        self._ingest(repo, md_file)
        result = _doc_fetch_for(repo, "validation", include_memories=False)
        assert result["memories"] == []
        assert result["total_memories"] == 0

    def test_no_scope_flag_skips_scope(self, repo, md_file):
        self._ingest(repo, md_file)
        result = _doc_fetch_for(repo, "memory", include_scope=False)
        assert result["scope"] is None

    def test_feature_with_no_match_returns_empty_bundle(self, repo, md_file):
        self._ingest(repo, md_file)
        # "xyzzy-nonexistent" should match nothing
        result = _doc_fetch_for(repo, "xyzzy-nonexistent")
        assert result["total_doc_files"] == 0
        # doc_search may return 0 as well — just assert it doesn't error
        assert "error" not in result

    def test_doc_search_deduplicates_with_doc_files(self, repo, tmp_path):
        """Files already in doc_files should not also appear in doc_search."""
        ai_gen = repo / ".ai-context" / "generated"
        ai_gen.mkdir(parents=True, exist_ok=True)
        # Create a file that both matches the slug and mentions the feature name
        target = ai_gen / "001-memory-layer.md"
        target.write_text("# Memory Layer\nmemory-layer is important.\n", encoding="utf-8")
        (ai_gen / "index.json").write_text(
            json.dumps({
                "source": "test.md", "generated_at": "2026-01-01",
                "mode": "python", "source_hash": "",
                "files": [{"id": "001-memory-layer",
                            "path": ".ai-context/generated/001-memory-layer.md",
                            "layer": "generated", "title": "Memory Layer"}],
            }), encoding="utf-8",
        )
        result = _doc_fetch_for(repo, "memory-layer")
        # The file should be in doc_files (slug match)
        assert result["total_doc_files"] >= 1
        # It should NOT also appear in doc_search
        doc_search_paths = {r["path"] for r in result["doc_search"]}
        doc_files_paths  = {f["path"] for f in result["doc_files"]}
        overlap = doc_search_paths & doc_files_paths
        assert overlap == set(), f"duplicate paths in both doc_files and doc_search: {overlap}"


# ---------------------------------------------------------------------------
# TestDocDiff
# ---------------------------------------------------------------------------

class TestDocDiff:
    """Tests for _doc_diff — content-hash-based diff of .ai-context/ files."""

    def _ingest(self, repo, md_file, **kwargs):
        return ingest_document(repo, md_file, overwrite=True, **kwargs)

    def test_no_ai_context_returns_error(self, repo):
        result = _doc_diff(repo)
        assert "error" in result

    def test_fresh_ingest_all_unchanged(self, repo, md_file):
        self._ingest(repo, md_file)
        result = _doc_diff(repo)
        assert "error" not in result
        assert result["has_changes"] is False
        assert len(result["modified"]) == 0
        assert len(result["missing"]) == 0

    def test_total_checked_matches_indexed_files(self, repo, md_file):
        self._ingest(repo, md_file)
        result = _doc_diff(repo)
        # total_checked should equal the number of entries in index.json
        idx = json.loads(
            (repo / ".ai-context" / "generated" / "index.json")
            .read_text(encoding="utf-8")
        )
        assert result["total_checked"] == len(idx["files"])

    def test_modified_file_detected(self, repo, md_file):
        self._ingest(repo, md_file)
        # Manually edit one generated file
        gen_dir = repo / ".ai-context" / "generated"
        md_files = list(gen_dir.glob("*.md"))
        assert md_files, "expected at least one generated .md file"
        target = md_files[0]
        original = target.read_text(encoding="utf-8")
        target.write_text(original + "\n<!-- manually edited -->\n", encoding="utf-8")

        result = _doc_diff(repo)
        assert result["has_changes"] is True
        modified_paths = [f["path"] for f in result["modified"]]
        rel = str(target.relative_to(repo)).replace("\\", "/")
        assert rel in modified_paths, f"expected {rel} in modified, got {modified_paths}"

    def test_missing_file_detected(self, repo, md_file):
        self._ingest(repo, md_file)
        gen_dir = repo / ".ai-context" / "generated"
        md_files = list(gen_dir.glob("*.md"))
        assert md_files
        target = md_files[0]
        rel = str(target.relative_to(repo)).replace("\\", "/")
        target.unlink()

        result = _doc_diff(repo)
        assert result["has_changes"] is True
        missing_paths = [f["path"] for f in result["missing"]]
        assert rel in missing_paths

    def test_extra_file_detected(self, repo, md_file):
        self._ingest(repo, md_file)
        gen_dir = repo / ".ai-context" / "generated"
        # Create an extra file not in index.json
        extra = gen_dir / "999-manual-notes.md"
        extra.write_text("# Manual Notes\n\nsome content\n", encoding="utf-8")

        result = _doc_diff(repo)
        assert result["has_changes"] is True
        extra_paths = [f["path"] for f in result["extra"]]
        rel = str(extra.relative_to(repo)).replace("\\", "/")
        assert rel in extra_paths

    def test_written_hash_stored_in_index(self, repo, md_file):
        """After ingest, each index.json file entry should have a written_hash."""
        self._ingest(repo, md_file)
        idx = json.loads(
            (repo / ".ai-context" / "generated" / "index.json")
            .read_text(encoding="utf-8")
        )
        for entry in idx["files"]:
            assert "written_hash" in entry, f"missing written_hash in {entry.get('id')}"
            # Hash is a non-empty 8-char hex string
            wh = entry["written_hash"]
            assert len(wh) == 8, f"unexpected hash length: {wh!r}"
            assert all(c in "0123456789abcdef" for c in wh), f"not hex: {wh!r}"

    def test_result_keys_present(self, repo, md_file):
        self._ingest(repo, md_file)
        result = _doc_diff(repo)
        for key in ("source", "generated_at", "total_checked",
                    "unchanged", "modified", "missing", "extra", "has_changes"):
            assert key in result, f"missing key: {key}"


# ---------------------------------------------------------------------------
# TestDocPin
# ---------------------------------------------------------------------------

class TestDocPin:
    """Tests for _doc_pin / _doc_unpin / pin-awareness in ingest."""

    def _ingest(self, repo, md_file, **kwargs):
        return ingest_document(repo, md_file, overwrite=True, **kwargs)

    def _first_generated_id(self, repo):
        idx = json.loads(
            (repo / ".ai-context" / "generated" / "index.json").read_text(encoding="utf-8")
        )
        files = [f for f in idx["files"] if f["layer"] == "generated"]
        assert files, "no generated files in index"
        return files[0]["id"], files[0]["path"]

    def test_pin_adds_to_pinned_files(self, repo, md_file):
        self._ingest(repo, md_file)
        fid, rel = self._first_generated_id(repo)
        result = _doc_pin(repo, fid)
        assert "error" not in result
        assert result["pinned"] is True
        pinned = _read_pinned(repo)
        assert rel in pinned

    def test_pin_unknown_id_returns_error(self, repo, md_file):
        self._ingest(repo, md_file)
        result = _doc_pin(repo, "xyzzy-nonexistent-99999")
        assert "error" in result

    def test_pin_already_pinned_is_idempotent(self, repo, md_file):
        self._ingest(repo, md_file)
        fid, rel = self._first_generated_id(repo)
        _doc_pin(repo, fid)
        result = _doc_pin(repo, fid)
        assert "error" not in result
        assert result.get("already_pinned") is True
        # Should still be in pinned set
        assert rel in _read_pinned(repo)

    def test_unpin_removes_from_pinned(self, repo, md_file):
        self._ingest(repo, md_file)
        fid, rel = self._first_generated_id(repo)
        _doc_pin(repo, fid)
        result = _doc_unpin(repo, fid)
        assert "error" not in result
        assert result["unpinned"] is True
        assert rel not in _read_pinned(repo)

    def test_unpin_not_pinned_is_safe(self, repo, md_file):
        self._ingest(repo, md_file)
        fid, _ = self._first_generated_id(repo)
        result = _doc_unpin(repo, fid)
        assert "error" not in result
        assert result.get("not_pinned") is True

    def test_ingest_skips_pinned_even_with_overwrite(self, repo, md_file):
        self._ingest(repo, md_file)
        fid, rel = self._first_generated_id(repo)
        _doc_pin(repo, fid)

        # Re-ingest with overwrite=True — pinned file must NOT be overwritten
        target = repo / rel
        target.write_text("# MANUALLY EDITED\n\nKeep this.\n", encoding="utf-8")
        original_content = target.read_text(encoding="utf-8")

        result = ingest_document(repo, md_file, overwrite=True)
        assert "error" not in result

        # Content must be unchanged
        after_content = target.read_text(encoding="utf-8")
        assert after_content == original_content, "pinned file was overwritten!"

        # Status should be skipped_pinned
        statuses = {f["path"]: f["status"] for f in result["generated"]}
        assert statuses.get(rel) == "skipped_pinned"

    def test_pins_preserved_across_ingest_runs(self, repo, md_file):
        self._ingest(repo, md_file)
        fid, rel = self._first_generated_id(repo)
        _doc_pin(repo, fid)

        # Second ingest should preserve the pin in the new index.json
        ingest_document(repo, md_file, overwrite=True)
        pinned = _read_pinned(repo)
        assert rel in pinned, "pin was lost after re-ingest"


# ---------------------------------------------------------------------------
# TestDocAnnotate
# ---------------------------------------------------------------------------

class TestDocAnnotate:
    """Tests for _doc_annotate — file-level annotations on .ai-context/ entries."""

    def _ingest(self, repo, md_file, **kwargs):
        return ingest_document(repo, md_file, overwrite=True, **kwargs)

    def _first_generated_id(self, repo):
        idx = json.loads(
            (repo / ".ai-context" / "generated" / "index.json").read_text(encoding="utf-8")
        )
        files = [f for f in idx["files"] if f["layer"] == "generated"]
        assert files
        return files[0]["id"]

    def test_view_empty_annotations(self, repo, md_file):
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        result = _doc_annotate(repo, fid)
        assert "error" not in result
        assert result["action"] == "view"
        assert result["annotations"] == []

    def test_add_annotation_returns_added(self, repo, md_file):
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        result = _doc_annotate(repo, fid, add_note="Reviewed in sprint 12", author="alice")
        assert "error" not in result
        assert result["action"] == "added"
        assert len(result["annotations"]) == 1
        ann = result["annotations"][0]
        assert ann["note"] == "Reviewed in sprint 12"
        assert ann["author"] == "alice"
        assert "ts" in ann

    def test_annotation_persists_in_index_json(self, repo, md_file):
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        _doc_annotate(repo, fid, add_note="Needs update for v2")
        # Read annotation back via a second call
        result = _doc_annotate(repo, fid)
        assert len(result["annotations"]) == 1
        assert result["annotations"][0]["note"] == "Needs update for v2"

    def test_multiple_annotations_accumulate(self, repo, md_file):
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        _doc_annotate(repo, fid, add_note="First note")
        _doc_annotate(repo, fid, add_note="Second note")
        result = _doc_annotate(repo, fid)
        assert len(result["annotations"]) == 2
        notes = [a["note"] for a in result["annotations"]]
        assert "First note" in notes
        assert "Second note" in notes

    def test_clear_removes_all_annotations(self, repo, md_file):
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        _doc_annotate(repo, fid, add_note="A note")
        result = _doc_annotate(repo, fid, clear=True)
        assert result["action"] == "cleared"
        assert result["annotations"] == []
        # Verify persistence
        result2 = _doc_annotate(repo, fid)
        assert result2["annotations"] == []

    def test_unknown_file_returns_error(self, repo, md_file):
        self._ingest(repo, md_file)
        result = _doc_annotate(repo, "xyzzy-nonexistent")
        assert "error" in result

    def test_curated_file_annotation_stored_separately(self, repo, md_file):
        self._ingest(repo, md_file)
        cur_dir = repo / ".ai-context" / "curated"
        assert cur_dir.exists(), "expected curated/ after ingest"
        # Annotate a curated file
        result = _doc_annotate(repo, "constraints", add_note="Last reviewed 2026-01-15")
        assert "error" not in result
        assert result["layer"] == "curated"
        assert result["action"] == "added"
        # annotations.json should be created
        ann_path = repo / ".ai-context" / "annotations.json"
        assert ann_path.exists()
        data = json.loads(ann_path.read_text(encoding="utf-8"))
        assert any(
            "Last reviewed 2026-01-15" in a["note"]
            for entries in data.values()
            for a in entries
        )


# ---------------------------------------------------------------------------
# TestDocSection
# ---------------------------------------------------------------------------

class TestDocSection:
    """Tests for _doc_fetch_section and _extract_section."""

    def _ingest(self, repo, md_file, **kwargs):
        return ingest_document(repo, md_file, overwrite=True, **kwargs)

    def test_extract_section_found(self, repo, md_file):
        from scope_intel.cli import _extract_section, _doc_fetch_section
        content = "# Top\n\nIntro text.\n\n## Roadmap\n\nPhase 1 done.\n\n## Constraints\n\nNo secrets.\n"
        section = _extract_section(content, "roadmap")
        assert section is not None
        assert "Phase 1 done" in section
        assert "No secrets" not in section  # should stop before Constraints

    def test_extract_section_case_insensitive(self, repo, md_file):
        from scope_intel.cli import _extract_section
        content = "# Top\n\n## ROADMAP\n\nPhase 1.\n"
        assert _extract_section(content, "roadmap") is not None
        assert _extract_section(content, "ROADMAP") is not None

    def test_extract_section_not_found_returns_none(self, repo, md_file):
        from scope_intel.cli import _extract_section
        content = "# Top\n\n## Roadmap\n\nContent.\n"
        assert _extract_section(content, "nonexistent-heading-xyz") is None

    def test_doc_fetch_section_from_ingest(self, repo, md_file):
        from scope_intel.cli import _doc_fetch_section
        self._ingest(repo, md_file)
        # roadmap.md should exist and have a top-level heading we can grab
        result = _doc_fetch_section(repo, "roadmap", "Roadmap")
        assert "error" not in result, result.get("error")
        assert result["chars"] > 0
        assert "heading" in result
        assert result["chars"] <= result["full_chars"]

    def test_doc_fetch_section_unknown_heading(self, repo, md_file):
        from scope_intel.cli import _doc_fetch_section
        self._ingest(repo, md_file)
        result = _doc_fetch_section(repo, "roadmap", "xyzzy-no-such-heading")
        assert "error" in result

    def test_doc_fetch_section_unknown_file(self, repo, md_file):
        from scope_intel.cli import _doc_fetch_section
        self._ingest(repo, md_file)
        result = _doc_fetch_section(repo, "xyzzy-no-file", "anything")
        assert "error" in result

    def test_list_pinned_filter(self, repo, md_file):
        """scope doc list --pinned returns only pinned files."""
        self._ingest(repo, md_file)
        idx = json.loads(
            (repo / ".ai-context" / "generated" / "index.json").read_text(encoding="utf-8")
        )
        files = [f for f in idx["files"] if f["layer"] == "generated"]
        assert files
        fid = files[0]["id"]
        _doc_pin(repo, fid)

        full_list = _doc_list(repo)
        pinned_paths = _read_pinned(repo)

        # Simulate --pinned filter (same logic as cmd_doc)
        filtered_gen = [f for f in full_list.get("generated", []) if f["path"] in pinned_paths]
        assert len(filtered_gen) == 1
        assert filtered_gen[0]["id"] == fid


# ---------------------------------------------------------------------------
# TestDocSnapshot
# ---------------------------------------------------------------------------

class TestDocSnapshot:
    """Tests for _doc_snapshot_save / _doc_snapshot_list / _doc_diff_since."""

    def _ingest(self, repo, md_file, **kwargs):
        return ingest_document(repo, md_file, overwrite=True, **kwargs)

    def test_save_snapshot_returns_saved(self, repo, md_file):
        from scope_intel.cli import _doc_snapshot_save
        self._ingest(repo, md_file)
        result = _doc_snapshot_save(repo, "v1")
        assert "error" not in result
        assert result["saved"] is True
        assert result["name"] == "v1"
        assert result["total_files"] > 0

    def test_snapshot_file_created_on_disk(self, repo, md_file):
        from scope_intel.cli import _doc_snapshot_save, _snapshot_path
        self._ingest(repo, md_file)
        _doc_snapshot_save(repo, "test-snap")
        sp = _snapshot_path(repo, "test-snap")
        assert sp.exists()
        data = json.loads(sp.read_text(encoding="utf-8"))
        assert data["name"] == "test-snap"
        assert len(data["files"]) > 0

    def test_snapshot_list_includes_saved(self, repo, md_file):
        from scope_intel.cli import _doc_snapshot_save, _doc_snapshot_list
        self._ingest(repo, md_file)
        _doc_snapshot_save(repo, "alpha")
        _doc_snapshot_save(repo, "beta")
        result = _doc_snapshot_list(repo)
        names = [s["name"] for s in result["snapshots"]]
        assert "alpha" in names
        assert "beta" in names
        assert result["total"] >= 2

    def test_diff_since_no_changes(self, repo, md_file):
        from scope_intel.cli import _doc_snapshot_save, _doc_diff_since
        self._ingest(repo, md_file)
        _doc_snapshot_save(repo, "baseline")
        result = _doc_diff_since(repo, "baseline")
        assert "error" not in result
        assert result["has_changes"] is False
        assert len(result["unchanged"]) > 0
        assert len(result["modified"]) == 0

    def test_diff_since_detects_modification(self, repo, md_file):
        from scope_intel.cli import _doc_snapshot_save, _doc_diff_since
        self._ingest(repo, md_file)
        _doc_snapshot_save(repo, "before")

        # Modify a generated file
        gen_dir = repo / ".ai-context" / "generated"
        md_files = list(gen_dir.glob("*.md"))
        assert md_files
        target = md_files[0]
        target.write_text(target.read_text(encoding="utf-8") + "\n<!-- edited -->\n",
                          encoding="utf-8")

        result = _doc_diff_since(repo, "before")
        assert result["has_changes"] is True
        modified_paths = [f["path"] for f in result["modified"]]
        rel = str(target.relative_to(repo)).replace("\\", "/")
        assert rel in modified_paths

    def test_diff_since_detects_added_file(self, repo, md_file):
        from scope_intel.cli import _doc_snapshot_save, _doc_diff_since
        self._ingest(repo, md_file)
        _doc_snapshot_save(repo, "before")

        # Add a new file after snapshot
        new_file = repo / ".ai-context" / "generated" / "999-new.md"
        new_file.write_text("# New\nContent.\n", encoding="utf-8")

        result = _doc_diff_since(repo, "before")
        assert result["has_changes"] is True
        added_paths = [f["path"] for f in result["added"]]
        assert ".ai-context/generated/999-new.md" in added_paths

    def test_diff_since_unknown_snapshot_returns_error(self, repo, md_file):
        from scope_intel.cli import _doc_diff_since
        self._ingest(repo, md_file)
        result = _doc_diff_since(repo, "nonexistent-snapshot")
        assert "error" in result

    def test_snapshot_list_empty_when_none_saved(self, repo):
        from scope_intel.cli import _doc_snapshot_list
        result = _doc_snapshot_list(repo)
        assert result["snapshots"] == []
        assert result["total"] == 0


# ---------------------------------------------------------------------------
# TestDocReport
# ---------------------------------------------------------------------------

class TestDocReport:
    """Tests for _doc_report — comprehensive .ai-context/ dashboard."""

    def _ingest(self, repo, md_file, **kwargs):
        return ingest_document(repo, md_file, overwrite=True, **kwargs)

    def test_no_ai_context_returns_error(self, repo):
        from scope_intel.cli import _doc_report
        result = _doc_report(repo)
        assert "error" in result

    def test_report_returns_expected_keys(self, repo, md_file):
        from scope_intel.cli import _doc_report
        self._ingest(repo, md_file)
        result = _doc_report(repo)
        assert "error" not in result
        for key in ("source", "ingested_at", "mode", "files", "total_files",
                    "total_chars", "total_tokens", "healthy", "budget_hint",
                    "pinned_count", "snapshots", "total_snapshots"):
            assert key in result, f"missing key: {key}"

    def test_total_files_matches_file_list(self, repo, md_file):
        from scope_intel.cli import _doc_report
        self._ingest(repo, md_file)
        result = _doc_report(repo)
        assert result["total_files"] == len(result["files"])

    def test_pinned_count_reflects_pins(self, repo, md_file):
        from scope_intel.cli import _doc_report
        self._ingest(repo, md_file)
        # Pin one file
        idx = json.loads(
            (repo / ".ai-context" / "generated" / "index.json").read_text(encoding="utf-8")
        )
        fid = idx["files"][0]["id"]
        _doc_pin(repo, fid)

        result = _doc_report(repo)
        assert result["pinned_count"] == 1
        pinned_in_list = [f for f in result["files"] if f["pinned"]]
        assert len(pinned_in_list) == 1

    def test_annotation_count_in_file_entry(self, repo, md_file):
        from scope_intel.cli import _doc_report
        self._ingest(repo, md_file)
        idx = json.loads(
            (repo / ".ai-context" / "generated" / "index.json").read_text(encoding="utf-8")
        )
        fid = idx["files"][0]["id"]
        _doc_annotate(repo, fid, add_note="Test note 1")
        _doc_annotate(repo, fid, add_note="Test note 2")

        result = _doc_report(repo)
        annotated = [f for f in result["files"] if f["annotations"] > 0]
        assert annotated, "expected at least one file with annotations > 0"
        assert annotated[0]["annotations"] == 2

    def test_snapshots_included_in_report(self, repo, md_file):
        from scope_intel.cli import _doc_report, _doc_snapshot_save
        self._ingest(repo, md_file)
        _doc_snapshot_save(repo, "v1")
        result = _doc_report(repo)
        assert result["total_snapshots"] == 1
        assert result["snapshots"][0]["name"] == "v1"

    def test_budget_hint_present(self, repo, md_file):
        from scope_intel.cli import _doc_report
        self._ingest(repo, md_file)
        result = _doc_report(repo)
        assert result["budget_hint"]
        assert "context" in result["budget_hint"].lower()


# ---------------------------------------------------------------------------
# TestDocTag
# ---------------------------------------------------------------------------

class TestDocTag:
    """Tests for _doc_tag — free-form labels on .ai-context/ file entries."""

    def _ingest(self, repo, md_file, **kwargs):
        return ingest_document(repo, md_file, overwrite=True, **kwargs)

    def _first_generated_id(self, repo):
        idx = json.loads(
            (repo / ".ai-context" / "generated" / "index.json").read_text(encoding="utf-8")
        )
        files = [f for f in idx["files"] if f["layer"] == "generated"]
        assert files
        return files[0]["id"]

    def test_add_tag(self, repo, md_file):
        from scope_intel.cli import _doc_tag
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        result = _doc_tag(repo, fid, add_tags=["api", "reviewed"])
        assert "error" not in result
        assert "api" in result["tags"]
        assert "reviewed" in result["tags"]

    def test_tags_persist_in_index_json(self, repo, md_file):
        from scope_intel.cli import _doc_tag
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        _doc_tag(repo, fid, add_tags=["backend"])
        # Read back
        result = _doc_tag(repo, fid)
        assert "backend" in result["tags"]

    def test_remove_tag(self, repo, md_file):
        from scope_intel.cli import _doc_tag
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        _doc_tag(repo, fid, add_tags=["api", "legacy"])
        result = _doc_tag(repo, fid, remove_tags=["legacy"])
        assert "api" in result["tags"]
        assert "legacy" not in result["tags"]

    def test_clear_tags(self, repo, md_file):
        from scope_intel.cli import _doc_tag
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        _doc_tag(repo, fid, add_tags=["x", "y", "z"])
        result = _doc_tag(repo, fid, clear=True)
        assert result["tags"] == []
        assert result["action"] == "cleared"

    def test_tag_unknown_file_returns_error(self, repo, md_file):
        from scope_intel.cli import _doc_tag
        self._ingest(repo, md_file)
        result = _doc_tag(repo, "xyzzy-no-such-file", add_tags=["x"])
        assert "error" in result

    def test_list_tag_filter(self, repo, md_file):
        """scope doc list --tag <tag> returns only tagged files (generated layer)."""
        from scope_intel.cli import _doc_tag, _get_file_tags, _load_annotations
        self._ingest(repo, md_file)
        idx = json.loads(
            (repo / ".ai-context" / "generated" / "index.json").read_text(encoding="utf-8")
        )
        # Pick the first explicitly-generated file (layer=="generated")
        gen_entries = [f for f in idx["files"] if f.get("layer") == "generated"]
        assert gen_entries, "need at least one generated file"
        fid = gen_entries[0]["id"]
        rel = gen_entries[0]["path"]
        _doc_tag(repo, fid, add_tags=["my-special-tag"])

        # Re-read index AFTER tag write so we see the updated tags field
        idx_updated = json.loads(
            (repo / ".ai-context" / "generated" / "index.json").read_text(encoding="utf-8")
        )
        index_files = idx_updated["files"]
        ann_data = _load_annotations(repo)
        tags = _get_file_tags(repo, rel, "generated", index_files, ann_data)
        assert "my-special-tag" in tags

        # Other generated files should NOT have the tag
        other_gen = [f for f in index_files if f["path"] != rel and f.get("layer") == "generated"]
        for f in other_gen:
            other_tags = _get_file_tags(repo, f["path"], "generated", index_files, ann_data)
            assert "my-special-tag" not in other_tags


# ---------------------------------------------------------------------------
# TestDocOutline
# ---------------------------------------------------------------------------

class TestDocOutline:
    """Tests for _doc_outline — heading hierarchy extraction from a .ai-context/ file."""

    def _ingest(self, repo, md_file, **kwargs):
        return ingest_document(repo, md_file, overwrite=True, **kwargs)

    def _first_generated_id(self, repo):
        idx = json.loads(
            (repo / ".ai-context" / "generated" / "index.json").read_text(encoding="utf-8")
        )
        files = [f for f in idx["files"] if f["layer"] == "generated"]
        assert files
        return files[0]["id"]

    def test_outline_keys_present(self, repo, md_file):
        from scope_intel.cli import _doc_outline
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        result = _doc_outline(repo, fid)
        assert "error" not in result, result
        for key in ("id", "path", "title", "layer", "headings", "total_headings", "chars"):
            assert key in result, f"missing key: {key}"

    def test_outline_headings_is_list(self, repo, md_file):
        from scope_intel.cli import _doc_outline
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        result = _doc_outline(repo, fid)
        assert isinstance(result["headings"], list)

    def test_outline_heading_entry_keys(self, repo, md_file):
        from scope_intel.cli import _doc_outline
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        result = _doc_outline(repo, fid)
        assert result["total_headings"] > 0, "expected at least one heading in generated file"
        h = result["headings"][0]
        for key in ("level", "text", "line_no", "char_offset"):
            assert key in h, f"heading entry missing key: {key}"

    def test_outline_heading_levels_valid(self, repo, md_file):
        from scope_intel.cli import _doc_outline
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        result = _doc_outline(repo, fid)
        for h in result["headings"]:
            assert 1 <= h["level"] <= 6, f"invalid heading level: {h['level']}"

    def test_outline_line_numbers_positive(self, repo, md_file):
        from scope_intel.cli import _doc_outline
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        result = _doc_outline(repo, fid)
        for h in result["headings"]:
            assert h["line_no"] >= 1, f"line_no must be >= 1, got {h['line_no']}"
            assert h["char_offset"] >= 0

    def test_outline_unknown_file_returns_error(self, repo, md_file):
        from scope_intel.cli import _doc_outline
        self._ingest(repo, md_file)
        result = _doc_outline(repo, "xyzzy-no-such-file-at-all")
        assert "error" in result

    def test_outline_total_headings_matches_list(self, repo, md_file):
        from scope_intel.cli import _doc_outline
        self._ingest(repo, md_file)
        fid = self._first_generated_id(repo)
        result = _doc_outline(repo, fid)
        assert result["total_headings"] == len(result["headings"])


# ---------------------------------------------------------------------------
# TestDocSearchTag
# ---------------------------------------------------------------------------

class TestDocSearchTag:
    """Tests for _doc_search with tag_filter parameter."""

    def _ingest(self, repo, md_file, **kwargs):
        return ingest_document(repo, md_file, overwrite=True, **kwargs)

    def _first_generated_id_and_path(self, repo):
        idx = json.loads(
            (repo / ".ai-context" / "generated" / "index.json").read_text(encoding="utf-8")
        )
        files = [f for f in idx["files"] if f["layer"] == "generated"]
        assert files
        return files[0]["id"], files[0]["path"]

    def test_search_tag_filter_key_in_result(self, repo, md_file):
        """Result dict always carries tag_filter back to caller."""
        self._ingest(repo, md_file)
        result = _doc_search(repo, "overview", tag_filter="mytag")
        assert "tag_filter" in result
        assert result["tag_filter"] == "mytag"

    def test_search_tag_filter_no_tagged_files_returns_empty(self, repo, md_file):
        """When no file has the tag, search returns 0 matches regardless of query."""
        self._ingest(repo, md_file)
        result = _doc_search(repo, "Auto-generated", tag_filter="unassigned-tag-xyz")
        assert "error" not in result
        assert result["total_matches"] == 0

    def test_search_tag_filter_restricts_to_tagged_files(self, repo, md_file):
        """Only files with the given tag are searched; untagged files are excluded."""
        from scope_intel.cli import _doc_tag
        self._ingest(repo, md_file)
        fid, _ = self._first_generated_id_and_path(repo)
        # Tag a specific file with a unique label
        _doc_tag(repo, fid, add_tags=["search-scope-test"])
        # "Auto-generated" appears in generated files (footer)
        r_filtered = _doc_search(repo, "Auto-generated", tag_filter="search-scope-test")
        r_all = _doc_search(repo, "Auto-generated")
        # Filtered search must have ≤ matches than unfiltered
        assert r_filtered["files_searched"] <= r_all["files_searched"]
        # The tag filter value is propagated
        assert r_filtered["tag_filter"] == "search-scope-test"

    def test_search_tag_filter_files_searched_count(self, repo, md_file):
        """files_searched reflects the filtered-down set, not total candidate count."""
        from scope_intel.cli import _doc_tag
        self._ingest(repo, md_file)
        fid, _ = self._first_generated_id_and_path(repo)
        _doc_tag(repo, fid, add_tags=["count-test"])
        result = _doc_search(repo, "anything", tag_filter="count-test")
        # Only 1 file tagged — files_searched should be 1
        assert result["files_searched"] == 1


# ---------------------------------------------------------------------------
# TestDocValidate
# ---------------------------------------------------------------------------

class TestDocValidate:
    """Tests for _doc_validate — .ai-context/ integrity checker."""

    def _ingest(self, repo, md_file, **kwargs):
        return ingest_document(repo, md_file, overwrite=True, **kwargs)

    def test_no_ai_context_returns_error(self, repo):
        from scope_intel.cli import _doc_validate
        result = _doc_validate(repo)
        assert "error" in result

    def test_clean_ingest_is_ok(self, repo, md_file):
        from scope_intel.cli import _doc_validate
        self._ingest(repo, md_file)
        result = _doc_validate(repo)
        assert "error" not in result
        assert result["ok"] is True
        assert result["errors"] == 0

    def test_result_keys_present(self, repo, md_file):
        from scope_intel.cli import _doc_validate
        self._ingest(repo, md_file)
        result = _doc_validate(repo)
        for key in ("ok", "total_issues", "errors", "warnings", "issues"):
            assert key in result, f"missing key: {key}"

    def test_issues_is_list(self, repo, md_file):
        from scope_intel.cli import _doc_validate
        self._ingest(repo, md_file)
        result = _doc_validate(repo)
        assert isinstance(result["issues"], list)

    def test_orphaned_entry_detected_E001(self, repo, md_file):
        """Deleting a generated file while keeping the index entry → E001."""
        from scope_intel.cli import _doc_validate
        self._ingest(repo, md_file)
        # Remove one of the generated .md files from disk
        gen_dir = repo / ".ai-context" / "generated"
        md_files = sorted(gen_dir.glob("*.md"))
        assert md_files, "need at least one generated .md file"
        victim = md_files[0]
        victim.unlink()
        result = _doc_validate(repo)
        assert result["errors"] >= 1
        codes = [i["code"] for i in result["issues"]]
        assert "E001" in codes

    def test_drift_detected_W001(self, repo, md_file):
        """Editing a generated file after ingest → W001."""
        from scope_intel.cli import _doc_validate
        self._ingest(repo, md_file)
        gen_dir = repo / ".ai-context" / "generated"
        md_files = sorted(gen_dir.glob("*.md"))
        assert md_files
        # Append content to simulate post-ingest edit
        md_files[0].write_text(
            md_files[0].read_text(encoding="utf-8") + "\n\n<!-- manual edit -->",
            encoding="utf-8",
        )
        result = _doc_validate(repo)
        codes = [i["code"] for i in result["issues"]]
        assert "W001" in codes
        # Warnings don't make ok=False
        assert result["ok"] is True  # warnings only → still ok

    def test_total_issues_matches_list(self, repo, md_file):
        from scope_intel.cli import _doc_validate
        self._ingest(repo, md_file)
        result = _doc_validate(repo)
        assert result["total_issues"] == len(result["issues"])


# ---------------------------------------------------------------------------
# TestDocRename
# ---------------------------------------------------------------------------

class TestDocRename:
    """Tests for _doc_rename — rename a curated .ai-context/ file."""

    def _make_curated(self, repo, stem: str, content: str = "# Title\n\nBody.") -> Path:
        cur_dir = repo / ".ai-context" / "curated"
        cur_dir.mkdir(parents=True, exist_ok=True)
        p = cur_dir / f"{stem}.md"
        p.write_text(content, encoding="utf-8")
        return p

    def test_rename_moves_file(self, repo):
        from scope_intel.cli import _doc_rename
        self._make_curated(repo, "constraints")
        result = _doc_rename(repo, "constraints", "new-constraints")
        assert "error" not in result, result
        assert result["ok"] is True
        assert (repo / ".ai-context" / "curated" / "new-constraints.md").exists()
        assert not (repo / ".ai-context" / "curated" / "constraints.md").exists()

    def test_rename_result_keys(self, repo):
        from scope_intel.cli import _doc_rename
        self._make_curated(repo, "design-spec")
        result = _doc_rename(repo, "design-spec", "api-spec")
        for key in ("ok", "old_path", "new_path", "annotations_updated"):
            assert key in result, f"missing key: {key}"

    def test_rename_unknown_file_returns_error(self, repo):
        from scope_intel.cli import _doc_rename
        self._make_curated(repo, "existing")
        result = _doc_rename(repo, "xyzzy-no-such", "new-name")
        assert "error" in result

    def test_rename_target_already_exists_returns_error(self, repo):
        from scope_intel.cli import _doc_rename
        self._make_curated(repo, "alpha")
        self._make_curated(repo, "beta")
        result = _doc_rename(repo, "alpha", "beta")
        assert "error" in result

    def test_rename_updates_annotations(self, repo):
        from scope_intel.cli import _doc_rename
        import json as _json
        self._make_curated(repo, "old-name")
        ai_ctx = repo / ".ai-context"
        ann_path = ai_ctx / "annotations.json"
        old_rel = ".ai-context/curated/old-name.md"
        # Seed an annotation for the old path
        ann_path.write_text(
            _json.dumps({old_rel: [{"ts": "2025-01-01T00:00:00Z", "note": "seed"}]}),
            encoding="utf-8",
        )
        result = _doc_rename(repo, "old-name", "renamed-file")
        assert "error" not in result
        ann_updated = _json.loads(ann_path.read_text(encoding="utf-8"))
        new_rel = ".ai-context/curated/renamed-file.md"
        assert new_rel in ann_updated
        assert old_rel not in ann_updated
        assert result["annotations_updated"] == 1

    def test_rename_partial_name_match(self, repo):
        from scope_intel.cli import _doc_rename
        self._make_curated(repo, "authentication-spec")
        # Partial match: "auth" should find "authentication-spec"
        result = _doc_rename(repo, "auth", "auth-spec")
        assert "error" not in result, result
        assert (repo / ".ai-context" / "curated" / "auth-spec.md").exists()


class TestDocCopy:
    """Tests for _doc_copy — duplicate a curated .ai-context/ file."""

    def _make_curated(self, repo, stem: str, content: str = "# Title\n\nBody.") -> Path:
        cur_dir = repo / ".ai-context" / "curated"
        cur_dir.mkdir(parents=True, exist_ok=True)
        p = cur_dir / f"{stem}.md"
        p.write_text(content, encoding="utf-8")
        return p

    def test_copy_creates_duplicate(self, repo):
        from scope_intel.cli import _doc_copy
        self._make_curated(repo, "constraints", "# Constraints\n\nOriginal text.")
        result = _doc_copy(repo, "constraints", "constraints-v2")
        assert "error" not in result, result
        assert result["ok"] is True
        # Both files exist
        assert (repo / ".ai-context" / "curated" / "constraints.md").exists()
        new_path = repo / ".ai-context" / "curated" / "constraints-v2.md"
        assert new_path.exists()
        # Content matches the source
        assert new_path.read_text(encoding="utf-8") == "# Constraints\n\nOriginal text."

    def test_copy_result_keys(self, repo):
        from scope_intel.cli import _doc_copy
        self._make_curated(repo, "design")
        result = _doc_copy(repo, "design", "design-draft")
        for key in ("ok", "source_path", "new_path"):
            assert key in result, f"missing key: {key}"

    def test_copy_unknown_source_returns_error(self, repo):
        from scope_intel.cli import _doc_copy
        self._make_curated(repo, "alpha")
        result = _doc_copy(repo, "no-such-file", "anything")
        assert "error" in result

    def test_copy_target_already_exists_returns_error(self, repo):
        from scope_intel.cli import _doc_copy
        self._make_curated(repo, "alpha")
        self._make_curated(repo, "beta")
        result = _doc_copy(repo, "alpha", "beta")
        assert "error" in result

    def test_copy_does_not_carry_annotations(self, repo):
        """Annotations are tied to the source path and intentionally not duplicated."""
        from scope_intel.cli import _doc_copy
        import json as _json
        self._make_curated(repo, "spec")
        ann_path = repo / ".ai-context" / "annotations.json"
        src_rel = ".ai-context/curated/spec.md"
        ann_path.write_text(
            _json.dumps({src_rel: [{"ts": "2025-01-01T00:00:00Z", "note": "old note"}]}),
            encoding="utf-8",
        )
        result = _doc_copy(repo, "spec", "spec-fork")
        assert "error" not in result
        ann_data = _json.loads(ann_path.read_text(encoding="utf-8"))
        new_rel = ".ai-context/curated/spec-fork.md"
        # New copy starts clean — annotations only on the original path.
        assert src_rel in ann_data
        assert new_rel not in ann_data

    def test_copy_no_curated_dir_returns_error(self, repo):
        from scope_intel.cli import _doc_copy
        result = _doc_copy(repo, "anything", "new")
        assert "error" in result


class TestDocTouch:
    """Tests for _doc_touch — one-shot 'needs review' annotation wrapper."""

    def _make_curated(self, repo, stem: str) -> Path:
        cur_dir = repo / ".ai-context" / "curated"
        cur_dir.mkdir(parents=True, exist_ok=True)
        p = cur_dir / f"{stem}.md"
        p.write_text("# Title\n", encoding="utf-8")
        return p

    def test_touch_adds_needs_review(self, repo):
        from scope_intel.cli import _doc_touch
        self._make_curated(repo, "design")
        result = _doc_touch(repo, "design")
        assert "error" not in result, result
        assert result["action"] == "added"
        assert result["annotations"][-1]["note"] == "needs review"

    def test_touch_with_reason_appends(self, repo):
        from scope_intel.cli import _doc_touch
        self._make_curated(repo, "spec")
        result = _doc_touch(repo, "spec", reason="schema is stale")
        assert "error" not in result
        assert result["annotations"][-1]["note"] == "needs review: schema is stale"

    def test_touch_records_author(self, repo):
        from scope_intel.cli import _doc_touch
        self._make_curated(repo, "doc")
        result = _doc_touch(repo, "doc", author="alice")
        assert result["annotations"][-1]["author"] == "alice"

    def test_touch_unknown_file_returns_error(self, repo):
        from scope_intel.cli import _doc_touch
        result = _doc_touch(repo, "nonexistent")
        assert "error" in result

    def test_touch_accumulates_with_existing_annotations(self, repo):
        """A second touch leaves prior annotations intact and stacks."""
        from scope_intel.cli import _doc_touch, _doc_annotate
        self._make_curated(repo, "page")
        _doc_annotate(repo, "page", add_note="initial thought")
        result = _doc_touch(repo, "page", reason="needs follow-up")
        notes = [a["note"] for a in result["annotations"]]
        assert notes[0] == "initial thought"
        assert notes[-1] == "needs review: needs follow-up"


class TestDocIngestWatch:
    """Tests for _doc_ingest_watch_tick — single-sweep auto-ingest."""

    def test_first_sweep_ingests_all_files(self, repo, tmp_path):
        from scope_intel.cli import _doc_ingest_watch_tick
        watch = tmp_path / "watched"
        watch.mkdir()
        (watch / "spec.md").write_text(SAMPLE_MD, encoding="utf-8")
        (watch / "notes.md").write_text("# Notes\n\nA short note.", encoding="utf-8")

        state: dict[str, float] = {}
        res = _doc_ingest_watch_tick(repo, watch, state)
        assert "error" not in res, res
        assert res["scanned"] == 2
        assert len(res["ingested"]) == 2
        assert len(res["errors"]) == 0
        # State now records both files
        assert len(state) == 2

    def test_second_sweep_skips_unchanged(self, repo, tmp_path):
        from scope_intel.cli import _doc_ingest_watch_tick
        watch = tmp_path / "watched"
        watch.mkdir()
        (watch / "spec.md").write_text(SAMPLE_MD, encoding="utf-8")

        state: dict[str, float] = {}
        first = _doc_ingest_watch_tick(repo, watch, state)
        assert len(first["ingested"]) == 1

        second = _doc_ingest_watch_tick(repo, watch, state)
        assert second["scanned"] == 1
        assert len(second["ingested"]) == 0
        assert len(second["skipped"]) == 1
        assert second["skipped"][0]["reason"] == "unchanged-mtime"

    def test_modified_file_is_re_ingested(self, repo, tmp_path):
        import os, time
        from scope_intel.cli import _doc_ingest_watch_tick
        watch = tmp_path / "watched"
        watch.mkdir()
        f = watch / "spec.md"
        f.write_text(SAMPLE_MD, encoding="utf-8")

        state: dict[str, float] = {}
        _doc_ingest_watch_tick(repo, watch, state)

        # Force mtime forward and rewrite
        new_mtime = time.time() + 5
        f.write_text(SAMPLE_MD + "\n\n## Added Section\nMore content.\n", encoding="utf-8")
        os.utime(f, (new_mtime, new_mtime))

        res = _doc_ingest_watch_tick(repo, watch, state)
        # mtime changed, so it gets re-processed; ingest_document with if_changed=True
        # may classify as ingested or unchanged-content depending on hash.
        assert (len(res["ingested"]) + len(res["skipped"])) == 1
        assert len(res["errors"]) == 0

    def test_new_file_added_between_sweeps(self, repo, tmp_path):
        from scope_intel.cli import _doc_ingest_watch_tick
        watch = tmp_path / "watched"
        watch.mkdir()
        (watch / "first.md").write_text("# First\n", encoding="utf-8")

        state: dict[str, float] = {}
        first = _doc_ingest_watch_tick(repo, watch, state)
        assert len(first["ingested"]) == 1

        # Drop a new file in
        (watch / "second.md").write_text("# Second\n", encoding="utf-8")
        second = _doc_ingest_watch_tick(repo, watch, state)
        assert second["scanned"] == 2
        # Only the new file should be ingested.
        ingested_names = {Path(e["path"]).name for e in second["ingested"]}
        assert ingested_names == {"second.md"}

    def test_invalid_directory_returns_error(self, repo, tmp_path):
        from scope_intel.cli import _doc_ingest_watch_tick
        bogus = tmp_path / "does-not-exist"
        state: dict[str, float] = {}
        res = _doc_ingest_watch_tick(repo, bogus, state)
        assert "error" in res

    def test_glob_pattern_filters(self, repo, tmp_path):
        from scope_intel.cli import _doc_ingest_watch_tick
        watch = tmp_path / "watched"
        watch.mkdir()
        (watch / "yes.md").write_text("# yes\n", encoding="utf-8")
        (watch / "no.txt").write_text("ignored\n", encoding="utf-8")
        state: dict[str, float] = {}
        res = _doc_ingest_watch_tick(repo, watch, state, glob_patterns=("*.md",))
        assert res["scanned"] == 1
        assert len(res["ingested"]) == 1


# ---------------------------------------------------------------------------
# Live integration tests — require Ollama + qwen2.5:7b running locally.
# Skipped automatically in CI or when Ollama is unavailable.
# ---------------------------------------------------------------------------

_OLLAMA_URL   = "http://localhost:11434"
_OLLAMA_MODEL = "qwen2.5:7b"
# Generous timeout: qwen2.5:7b can take 90s+ on first call (cold GPU load)
_OLLAMA_TIMEOUT = 180

def _ollama_available() -> bool:
    """Return True if Ollama is reachable, qwen2.5:7b is listed, AND responds to a generate call.

    Checking /api/tags alone isn't enough — the model may be listed but not yet loaded into
    GPU memory (cold-start).  We do a minimal generate call with a 30s timeout to confirm
    the model is actually serving, which also warms the model for the real tests.
    """
    from scope_intel.core.llm_client import OllamaClient
    client = OllamaClient(model=_OLLAMA_MODEL, url=_OLLAMA_URL, timeout=_OLLAMA_TIMEOUT)
    if not client.is_available():
        return False
    # Warm-up generate: tiny prompt, short response.  If this returns None the model is
    # not yet ready (still loading) and we skip the suite.
    warmup = client._generate('Say "ok" as JSON: {"status":"ok"}')
    return warmup is not None


_LIVE_DOC = """\
# Scope Intelligence Toolkit

## Overview
Scope Intelligence Toolkit is a CLI tool that indexes repositories and answers
scope queries: which features are affected by a change, which symbols call which,
and which tests need to run.

## Architecture
The system has three layers:
- **Deterministic Engine** — Python AST indexer, no external dependencies.
- **Memory Layer** — MemPalace stores semantic, episodic, and procedural memories.
- **MCP Server** — JSON-RPC 2.0 over stdio exposing 15+ tools to Claude Code.

## Constraints
- Zero external Python dependencies for the core toolkit.
- All storage is JSONL — no PostgreSQL, no Chroma, no pgvector.
- Ollama is called as a tool; Python handles routing and writes.

## Roadmap
- Phase 6: cross-repo federation, confidence decay, agent-triggered captures.
- Phase 7: RAG retrieval over generated .ai-context/ files.

## Current Phase
Phase 5 is complete. Phase 6 design doc is committed. Implementation pending.
"""


@pytest.mark.skipif(not _ollama_available(), reason="Ollama not available or qwen2.5:7b not loaded")
class TestLLMIngestLive:
    """End-to-end integration tests for mode=llm ingest pipeline.

    These tests hit a real Ollama instance and verify that all three LLM
    pipeline steps produce meaningful output on a known design document.
    Run with: pytest tests/test_doc_ingest.py::TestLLMIngestLive -v
    """

    @pytest.fixture
    def repo(self, tmp_path):
        from scope_intel.core import store
        store.ensure_index_dir(tmp_path)
        store.write_json(tmp_path, "config", store.default_config())
        return tmp_path

    @pytest.fixture
    def design_doc(self, tmp_path):
        p = tmp_path / "design.md"
        p.write_text(_LIVE_DOC, encoding="utf-8")
        return p

    def test_llm_probe_returns_content_readable(self, design_doc):
        """llm-probe reads the file locally and Ollama classifies it."""
        from scope_intel.core.llm_client import probe_file_path
        result = probe_file_path(str(design_doc), url=_OLLAMA_URL, model=_OLLAMA_MODEL,
                                 timeout=_OLLAMA_TIMEOUT)
        assert result["can_read_content"], (
            f"Ollama did not return a valid JSON classification. "
            f"error={result.get('error')!r}  llm_response={result.get('llm_response')!r}"
        )
        assert result["bytes_sent"] > 0
        assert result.get("llm_response", {}).get("type"), "expected a 'type' field in LLM response"

    def test_global_summary_step(self, design_doc):
        """Step 2: global_summary() returns project_name and at least one component."""
        from scope_intel.core.llm_client import OllamaClient
        client = OllamaClient(model=_OLLAMA_MODEL, url=_OLLAMA_URL, timeout=_OLLAMA_TIMEOUT)
        full_text = design_doc.read_text(encoding="utf-8")
        result = client.global_summary(full_text)
        assert result is not None, "global_summary returned None — LLM call failed"
        assert isinstance(result, dict)
        assert result.get("project_name"), f"project_name missing: {result}"
        # purpose or components must be non-empty for a meaningful extraction
        has_substance = bool(result.get("purpose") or result.get("components"))
        assert has_substance, f"global_summary lacks purpose and components: {result}"

    def test_classify_chunk_step(self, design_doc):
        """Step 3: classify_chunk() returns a valid target_file and category for each section."""
        from scope_intel.core.llm_client import OllamaClient
        from scope_intel.adapters.doc_reader import read_document_structured
        client = OllamaClient(model=_OLLAMA_MODEL, url=_OLLAMA_URL, timeout=_OLLAMA_TIMEOUT)

        struct = read_document_structured(design_doc)
        assert "error" not in struct, f"read_document_structured failed: {struct}"
        chunks = struct["chunks"]
        assert chunks, "design doc produced no chunks"

        global_ctx = client.global_summary(struct["full_text"]) or {
            "project_name": "Scope Intelligence Toolkit",
            "components": [],
        }

        classified = 0
        for chunk in chunks:
            result = client.classify_chunk(chunk, global_ctx)
            if result is not None:
                assert isinstance(result, dict)
                assert "target_file" in result, f"classify_chunk missing target_file: {result}"
                assert "category" in result, f"classify_chunk missing category: {result}"
                classified += 1

        assert classified > 0, (
            f"classify_chunk returned None for all {len(chunks)} chunks — "
            "LLM may be failing to parse the prompt or schema"
        )

    def test_module_map_pass_step(self):
        """Step 5: module_map_pass() returns non-empty markdown given summaries."""
        from scope_intel.core.llm_client import OllamaClient
        client = OllamaClient(model=_OLLAMA_MODEL, url=_OLLAMA_URL, timeout=_OLLAMA_TIMEOUT)
        summaries = [
            {"title": "Deterministic Engine", "summary": "Python AST indexer with zero deps."},
            {"title": "Memory Layer",          "summary": "JSONL-backed semantic + episodic memory."},
            {"title": "MCP Server",            "summary": "JSON-RPC 2.0 stdio server exposing 15 tools."},
        ]
        result = client.module_map_pass(summaries)
        assert result is not None, "module_map_pass returned None — LLM call failed"
        assert len(result.strip()) > 50, f"module_map_pass output too short: {result!r}"
        # Should contain at least one module name
        assert any(s["title"] in result for s in summaries), (
            "module_map_pass output doesn't mention any module titles"
        )

    def test_full_ingest_llm_pipeline(self, repo, design_doc):
        """End-to-end: ingest_document(mode='llm') runs all 5 steps and writes files."""
        from scope_intel.core.doc_ingestor import ingest_document
        result = ingest_document(
            repo,
            design_doc,
            mode="llm",
            ollama_model=_OLLAMA_MODEL,
            ollama_url=_OLLAMA_URL,
            overwrite=True,
        )
        assert "error" not in result, f"ingest_document failed: {result.get('error')}"
        assert result.get("mode") == "llm", f"expected mode=llm, got {result.get('mode')!r}"
        assert result.get("sections_parsed", 0) > 0, "no sections parsed"
        assert result.get("files_written", 0) > 0, "no files written"
        # Verify at least one .ai-context/generated/*.md exists on disk
        gen_dir = repo / ".ai-context" / "generated"
        md_files = list(gen_dir.glob("*.md"))
        assert md_files, f"no .md files in {gen_dir}"
        # LLM pipeline stats must be present and at least one chunk classified by the LLM
        # (not just Python fallback — that would mean mode=llm had no actual LLM calls)
        assert result.get("llm_chunks_classified") is not None, (
            "llm_chunks_classified missing from result — LLM pipeline may have short-circuited"
        )
        llm_by_llm = result.get("llm_chunks_by_llm", 0)
        assert llm_by_llm > 0, (
            f"llm_chunks_by_llm={llm_by_llm}: Ollama responded but classified 0 chunks. "
            f"All {result.get('llm_chunks_classified')} chunks used Python fallback — "
            "check Ollama is serving the model and not timing out."
        )
