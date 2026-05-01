"""Document readers — PDF, DOCX, MD, TXT.

All formats normalise to plain markdown-style text so the ingestor
sees a single consistent input regardless of file type.

Two reading modes:

  read_document(path)
      → {"text": str, "format": str, ...}
      Plain text. Used by mode=python (existing pipeline, no LLM).

  read_document_structured(path)
      → {"chunks": list[Chunk], "format": str, "total_chunks": int, ...}
      Splits by headings (any # level — generic, no keyword matching).
      Each chunk carries a heading_path so the LLM classifier knows
      where in the document the chunk came from even after splitting.

      Chunk schema:
        {
          "heading_path": "Architecture > Memory Layer > Redis Setup",
          "title":        "Redis Setup",
          "level":        3,
          "text":         "...",
          "token_estimate": 412,
          "chunk_index":  7,
        }

      Long sections (> MAX_CHUNK_TOKENS) are split with OVERLAP_TOKENS
      overlap so context is not lost at split boundaries.

Tables and images:
  MD  tables  — already plain text, pass through as-is.
  DOCX tables — extracted in document order, rendered as markdown tables.
  PDF  tables — pdfplumber (if installed) extracts bounding-box-aware tables
                as markdown; pypdf fallback provides text-only (may scramble cols).
  DOCX images — alt-text / title attributes extracted as [Image: description]
                markers (no extra deps — uses XML attributes Word writes by default).
  PDF images  — silently skipped (binary pixel data requires OCR or vision model).
  All images  — for OCR use pytesseract; for visual semantics use a vision model
                (e.g. qwen-vl / llava via Ollama /api/generate with images field).

Optional dependencies (installed only if you ingest that format):
  PDF (best)  → pip install pdfplumber   ← tables + reading order
  PDF (basic) → pip install pypdf        ← text only, no table structure
  DOCX        → pip install python-docx

MD / TXT need no extra packages.
"""
from __future__ import annotations

import re
from pathlib import Path

# Chunking parameters — tuned for Qwen2.5:14b context window
MAX_CHUNK_TOKENS = 1200   # split sections larger than this
OVERLAP_TOKENS   = 150    # token overlap between adjacent sub-chunks
_CHARS_PER_TOKEN = 4      # rough estimate: 1 token ≈ 4 chars (English prose)


def read_document(path: Path) -> dict:
    """Read a document file and return normalised text.

    Returns:
        {"text": str, "format": str, ...meta}  on success
        {"error": str}                          on failure
    """
    path = Path(path)
    if not path.exists():
        return {"error": f"file not found: {path}"}

    ext = path.suffix.lower()

    if ext in (".md", ".txt", ".rst", ".text"):
        return _read_text(path)
    if ext == ".pdf":
        return _read_pdf(path)
    if ext in (".docx",):
        return _read_docx(path)
    # unknown extension — try as plain text
    return _read_text(path)


# ---------------------------------------------------------------------------
# Plain text / Markdown
# ---------------------------------------------------------------------------

def _read_text(path: Path) -> dict:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        return {"text": text, "format": path.suffix.lstrip(".") or "txt"}
    except OSError as exc:
        return {"error": str(exc)}


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _read_pdf(path: Path) -> dict:
    """Read a PDF file.

    Strategy (best-effort, graceful degradation):
      1. Try pdfplumber first — understands bounding boxes, extracts
         tables as markdown and text in reading order. Optional dep.
      2. Fall back to pypdf — text only, no table structure.
         Cells from multi-column layouts may be interleaved.
      3. Fail with a clear message if neither is installed.

    Install:
      pip install pdfplumber   ← best (tables + reading order)
      pip install pypdf        ← fallback (text only)
    """
    try:
        import pdfplumber  # type: ignore
        return _read_pdf_pdfplumber(path, pdfplumber)
    except ImportError:
        pass  # fall through to pypdf

    try:
        import pypdf  # type: ignore
        return _read_pdf_pypdf(path, pypdf)
    except ImportError:
        return {
            "error": (
                "No PDF library found.\n"
                "Install pdfplumber (recommended):  pip install pdfplumber\n"
                "Or pypdf (text only):              pip install pypdf"
            )
        }


def _read_pdf_pdfplumber(path: Path, pdfplumber) -> dict:
    """Extract text + tables from PDF using pdfplumber (best quality)."""
    try:
        pages_text: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            num_pages = len(pdf.pages)
            for page in pdf.pages:
                parts: list[str] = []

                # Extract tables first, render as markdown
                tables_on_page = page.extract_tables() or []
                # Track bounding boxes of extracted tables so we can exclude
                # their text from the plain-text pass (avoid double-reading)
                table_bboxes: list[tuple] = []
                for table_obj in page.find_tables():
                    table_bboxes.append(table_obj.bbox)  # (x0, top, x1, bottom)

                # Page words not inside any table bbox
                page_words = page.extract_words() or []
                non_table_words = [
                    w for w in page_words
                    if not any(
                        w["x0"] >= bbox[0] and w["top"] >= bbox[1]
                        and w["x1"] <= bbox[2] and w["bottom"] <= bbox[3]
                        for bbox in table_bboxes
                    )
                ]
                # Re-assemble prose text from filtered words
                prose = " ".join(w["text"] for w in non_table_words).strip()
                if prose:
                    parts.append(prose)

                # Render each table as markdown
                for rows in tables_on_page:
                    if not rows:
                        continue
                    rows = [[cell or "" for cell in row] for row in rows]
                    header = rows[0]
                    md = (
                        "| " + " | ".join(header) + " |\n"
                        + "| " + " | ".join(["---"] * len(header)) + " |\n"
                        + "\n".join(
                            "| " + " | ".join(row[:len(header)]) + " |"
                            for row in rows[1:]
                        )
                    )
                    parts.append(md)

                if parts:
                    pages_text.append("\n\n".join(parts))

        if not pages_text:
            return {"error": "PDF produced no extractable text (may be image-only)"}
        return {
            "text": _normalise_pdf_text("\n\n".join(pages_text)),
            "format": "pdf",
            "pages": num_pages,
            "reader": "pdfplumber",
        }
    except Exception as exc:  # noqa: BLE001
        return {"error": f"PDF read error (pdfplumber): {exc}"}


def _read_pdf_pypdf(path: Path, pypdf) -> dict:
    """Extract text from PDF using pypdf (text only, no table structure)."""
    try:
        reader = pypdf.PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text)
        if not pages:
            return {"error": "PDF produced no extractable text (may be image-only)"}
        return {
            "text": _normalise_pdf_text("\n\n".join(pages)),
            "format": "pdf",
            "pages": len(reader.pages),
            "reader": "pypdf",
        }
    except Exception as exc:  # noqa: BLE001
        return {"error": f"PDF read error (pypdf): {exc}"}


def _normalise_pdf_text(text: str) -> str:
    """Best-effort cleanup: rejoin hyphenated line-breaks, normalise whitespace."""
    # Rejoin words broken at line end with a hyphen
    text = re.sub(r"-\n(\w)", r"\1", text)
    # Collapse runs of blank lines to a double newline
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _docx_table_to_markdown(table) -> str:
    """Convert a python-docx Table object to a GitHub-flavoured markdown table.

    Row 0 is treated as the header row.
    Cells with merged columns are repeated (best-effort).
    """
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    # De-duplicate adjacent identical cells (column-spanning merged cells)
    def _dedup(row: list[str]) -> list[str]:
        out: list[str] = []
        for cell in row:
            if not out or cell != out[-1]:
                out.append(cell)
        return out

    header = _dedup(rows[0])
    col_count = len(header)
    lines: list[str] = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * col_count) + " |",
    ]
    for row in rows[1:]:
        row = _dedup(row)
        # Pad / truncate to match header width
        while len(row) < col_count:
            row.append("")
        row = row[:col_count]
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


_WP_DRAWING_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_WP_DOCPR_TAG  = f"{{{_WP_DRAWING_NS}}}docPr"


def _extract_image_alt_texts(element) -> list[str]:
    """Return alt-text / title strings for any images embedded in this XML element.

    DOCX stores image alt-text in <wp:docPr descr="..." title="..."/> inside
    the paragraph that contains the drawing. No extra dependencies needed —
    uses the lxml etree that python-docx already requires.
    """
    labels: list[str] = []
    for doc_pr in element.iter(_WP_DOCPR_TAG):
        descr = (doc_pr.get("descr") or "").strip()
        title = (doc_pr.get("title") or "").strip()
        label = descr or title
        if label:
            labels.append(label)
    return labels


def _read_docx(path: Path) -> dict:
    """Read a DOCX file preserving document order: paragraphs, tables, and images.

    - Paragraphs: rendered as plain text (Heading styles → markdown #)
    - Tables: rendered as GitHub-flavoured markdown tables
    - Images: alt-text / title extracted as [Image: description] markers
      (no extra deps — uses DOCX XML attributes that Word writes by default)
    """
    try:
        import docx  # type: ignore  (python-docx)
    except ImportError:
        return {
            "error": (
                "python-docx is required to read DOCX files.\n"
                "Install it with:  pip install python-docx"
            )
        }
    try:
        doc = docx.Document(str(path))

        # Build lookup: XML element → paragraph/table object so we can
        # iterate body children in document order (paragraphs interleaved
        # with tables) rather than doc.paragraphs (paragraphs only).
        para_map  = {p._element: p for p in doc.paragraphs}
        table_map = {t._element: t for t in doc.tables}

        lines: list[str] = []
        for child in doc.element.body:
            if child in para_map:
                para = para_map[child]

                # 1. Image alt-text (may coexist with or replace paragraph text)
                for label in _extract_image_alt_texts(child):
                    lines.append(f"[Image: {label}]")

                # 2. Paragraph text
                text = para.text.strip()
                if not text:
                    continue
                style = para.style.name or ""
                if style.startswith("Heading"):
                    # "Heading 1" → #, "Heading 2" → ##, etc.
                    try:
                        level = int(style.split()[-1])
                    except ValueError:
                        level = 2
                    lines.append("#" * min(level, 4) + " " + text)
                elif style in ("Title",):
                    # Document Title style → h1 (the document heading)
                    lines.append("# " + text)
                elif style in ("Subtitle",):
                    # Document Subtitle → italic paragraph (not a heading,
                    # but worth preserving as context for the LLM)
                    lines.append(f"_{text}_")
                elif style in ("Caption",):
                    # Figure/table captions → preserve as italic note
                    lines.append(f"_{text}_")
                else:
                    lines.append(text)

            elif child in table_map:
                md_table = _docx_table_to_markdown(table_map[child])
                if md_table:
                    lines.append(md_table)

            # else: sectPr, bookmarks, etc. — skip silently

        if not lines:
            return {"error": "DOCX produced no text"}
        return {"text": "\n\n".join(lines), "format": "docx"}
    except Exception as exc:  # noqa: BLE001
        return {"error": f"DOCX read error: {exc}"}


# ---------------------------------------------------------------------------
# Structured reader — heading_path chunks (used by mode=llm)
# ---------------------------------------------------------------------------

def _estimate_tokens(text: str) -> int:
    return max(1, len(text) // _CHARS_PER_TOKEN)


def _split_by_tokens(text: str, max_tokens: int, overlap_tokens: int) -> list[str]:
    """Split a long text into overlapping sub-chunks.

    Splits at sentence boundaries when possible ('. ', '\\n') to avoid
    cutting mid-sentence. Each sub-chunk except the first starts with
    the last `overlap_tokens` worth of the previous sub-chunk.
    """
    if _estimate_tokens(text) <= max_tokens:
        return [text]

    max_chars    = max_tokens    * _CHARS_PER_TOKEN
    overlap_chars = overlap_tokens * _CHARS_PER_TOKEN

    chunks: list[str] = []
    pos = 0
    while pos < len(text):
        end = min(pos + max_chars, len(text))
        # Try to cut at a paragraph or sentence boundary
        if end < len(text):
            for sep in ("\n\n", "\n", ". ", " "):
                cut = text.rfind(sep, pos, end)
                if cut > pos:
                    end = cut + len(sep)
                    break
        chunks.append(text[pos:end].strip())
        if end >= len(text):
            break
        # Next chunk starts with overlap
        pos = max(pos + 1, end - overlap_chars)

    return [c for c in chunks if c]


def read_document_structured(path: Path) -> dict:
    """Read a document and return heading-aware chunks with heading_path.

    Heading split is GENERIC — splits on any #/##/###/#### pattern.
    No keyword matching here. The LLM classifier decides meaning.

    Each chunk:
      heading_path   — cumulative path: "Architecture > Memory > Redis"
      title          — this section's heading text
      level          — heading level (1-4)
      text           — section body (may be a sub-chunk of a long section)
      token_estimate — rough token count
      chunk_index    — global sequential index across all chunks

    Long sections are split with OVERLAP_TOKENS overlap so the
    LLM classifier never loses context at split boundaries.
    """
    # First get the plain text (reuse existing readers)
    read_result = read_document(path)
    if "error" in read_result:
        return read_result

    text: str = read_result["text"]
    fmt: str  = read_result.get("format", "txt")

    # ---- Parse into raw sections by heading ----
    raw_sections: list[dict] = []
    heading_stack: list[str] = []   # tracks current path per level
    current: dict = {"level": 0, "title": "(preamble)", "body_lines": []}

    for line in text.splitlines():
        m = re.match(r"^(#{1,4})\s+(.+)$", line)
        if m:
            # Flush current section
            body = "\n".join(current["body_lines"]).strip()
            if body or current["title"] != "(preamble)":
                raw_sections.append({
                    "level": current["level"],
                    "title": current["title"],
                    "body":  body,
                })
            # Update heading stack
            level = len(m.group(1))
            title = m.group(2).strip()
            # Trim stack to parent level
            heading_stack = heading_stack[: level - 1]
            heading_stack.append(title)
            current = {"level": level, "title": title, "body_lines": []}
        else:
            current["body_lines"].append(line)

    # Flush last section
    body = "\n".join(current["body_lines"]).strip()
    if body or current["title"] != "(preamble)":
        raw_sections.append({
            "level": current["level"],
            "title": current["title"],
            "body":  body,
        })

    # ---- Build chunks with heading_path ----
    chunks: list[dict] = []
    chunk_index = 0
    # Maintain a heading_path stack: level → title
    path_stack: list[tuple[int, str]] = []  # (level, title)

    for sec in raw_sections:
        lvl   = sec["level"]
        title = sec["title"]
        body  = sec["body"]

        # Rebuild path_stack to correct depth
        path_stack = [(l, t) for l, t in path_stack if l < lvl]
        path_stack.append((lvl, title))
        heading_path = " > ".join(t for _, t in path_stack)

        if not body:
            # Section with no body (just a heading) — still emit a minimal chunk
            # so the LLM can see the section exists
            chunks.append({
                "heading_path":   heading_path,
                "title":          title,
                "level":          lvl,
                "text":           f"[Section: {title}]",
                "token_estimate": 5,
                "chunk_index":    chunk_index,
            })
            chunk_index += 1
            continue

        # Split long sections to stay within LLM context window
        sub_texts = _split_by_tokens(body, MAX_CHUNK_TOKENS, OVERLAP_TOKENS)
        for i, sub in enumerate(sub_texts):
            sub_heading = heading_path if len(sub_texts) == 1 else f"{heading_path} (part {i+1}/{len(sub_texts)})"
            chunks.append({
                "heading_path":   sub_heading,
                "title":          title,
                "level":          lvl,
                "text":           sub,
                "token_estimate": _estimate_tokens(sub),
                "chunk_index":    chunk_index,
            })
            chunk_index += 1

    return {
        "chunks":        chunks,
        "total_chunks":  len(chunks),
        "format":        fmt,
        "source":        str(path),
        # Also include full text for global summary pass
        "full_text":     text,
    }
